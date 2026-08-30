# -*- coding: utf-8 -*-
"""
AgnesAgent v2 - 免费全能 AI 桌面助手
功能：对话问答 / 联网搜索（网页+图片，结果卡片化）/ 代码执行 / 文件读取 / 图像生成 / 图像理解 / 语音条回复
特性：
  - 多 Provider 自动兜底：主 API 不可用时自动切换到备用免费 API（智谱/百度/Ollama 本地）
  - 搜索结果卡片化：网页搜索结果、图片缩略图网格直接渲染在聊天区，不再空白
  - 语音条：TTS 生成后内嵌聊天气泡语音条，点击播放，不再弹出文件
  - 消息气泡 UI：仿聊天软件布局
"""
import sys
import os
import io
import json
import time
import random
import uuid
import base64
import hashlib
import threading
import subprocess
import requests

# ===== 开源库集成（GitHub 保底方案，全部防御式导入）=====
try:
    from g4f.client import Client as G4FClient  # gpt4free：零 Key 免费模型兜底
    G4F_AVAILABLE = True
except Exception:
    G4FClient = None
    G4F_AVAILABLE = False
try:
    from pypdf import PdfReader              # PDF 读取
except Exception:
    PdfReader = None
try:
    from docx import Document as DocxDocument  # Word 读取
except Exception:
    DocxDocument = None
try:
    from openpyxl import load_workbook        # Excel 读取
except Exception:
    load_workbook = None
try:
    from ddgs import DDGS                    # duckduckgo-search：免 Key 搜索保底
except Exception:
    DDGS = None

from PySide6.QtCore import Qt, QThread, Signal, QSize, QUrl, QTimer
from PySide6.QtGui import QFont, QIcon, QPixmap, QTextCursor, QColor, QDesktopServices, QTextOption
from PySide6.QtMultimedia import QMediaPlayer, QAudioOutput
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QTextBrowser,
    QTextEdit, QLineEdit, QPushButton, QLabel, QSplitter, QListWidget, QListWidgetItem,
    QDialog, QFormLayout, QComboBox, QSpinBox, QCheckBox, QFileDialog,
    QMessageBox, QScrollArea, QFrame, QGridLayout, QSizePolicy, QTableWidget,
    QTableWidgetItem, QHeaderView, QMenu, QGraphicsDropShadowEffect
)

# ============ Fluent 开源 UI（QFluentWidgets，Win11 Fluent 风格）============
# 引入失败时自动回退原生样式，不影响主功能
FLUENT_OK = False
try:
    from qfluentwidgets import (
        Theme, setTheme, setThemeColor, SwitchButton, PrimaryPushButton,
        CardWidget, BodyLabel, CaptionLabel, ToolTipFilter, InfoBadge,
        FluentIcon, IconWidget, PushButton, LineEdit, ProgressBar, AvatarWidget,
        ComboBox, Dialog, MessageBoxBase, TeachingTip, TeachingTipTailPosition,
    )
    FLUENT_OK = True
except Exception as _fqe:
    print("QFluentWidgets 不可用，回退原生样式:", _fqe)

# ===== 本地 AI 开源库（离线能力，全部防御式导入）=====
# 本地 OCR：RapidOCR（PaddleOCR 模型转 ONNX，完全离线）
try:
    from rapidocr_onnxruntime import RapidOCR as _RapidOCR
    RAPIDOCR_AVAILABLE = True
except Exception:
    _RapidOCR = None
    RAPIDOCR_AVAILABLE = False
# 本地 ASR/TTS：sherpa-onnx（k2-fsa 开源语音工具链，ONNX 推理）
try:
    import sherpa_onnx
    SHERPA_ONNX_AVAILABLE = True
except Exception:
    sherpa_onnx = None
    SHERPA_ONNX_AVAILABLE = False
# 本地 PDF 强解析：PyMuPDF（比 pypdf 更完整，支持复杂版式/表格/图片）
try:
    import pymupdf
    PYMUPDF_AVAILABLE = True
except Exception:
    try:
        import fitz as pymupdf
        PYMUPDF_AVAILABLE = True
    except Exception:
        pymupdf = None
        PYMUPDF_AVAILABLE = False

# 可写数据目录：打包(frozen)后始终为 exe 同级目录（config/history/images 持久化）；开发期为脚本同目录
if getattr(sys, "frozen", False):
    APP_DIR = os.path.dirname(sys.executable)
else:
    APP_DIR = os.path.dirname(os.path.abspath(__file__))

# 只读资源目录：onefile 打包时模型等资源解压在 _MEIPASS；onedir/开发期为 APP_DIR
if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
    RES_DIR = sys._MEIPASS
else:
    RES_DIR = APP_DIR

# 本地模型目录（打进 exe 的资源在 _MEIPASS/dl/；开发期为脚本同目录 dl/）
MODELS_DIR = os.path.join(RES_DIR, "dl", "models")
if not os.path.isdir(MODELS_DIR):
    MODELS_DIR = os.path.join(RES_DIR, "models")
LOCAL_ASR_MODEL = os.path.join(MODELS_DIR, "asr_sense_voice.onnx")
LOCAL_ASR_TOKENS = os.path.join(MODELS_DIR, "asr_tokens.txt")
LOCAL_TTS_MODEL = os.path.join(MODELS_DIR, "tts_vits_zh_aishell3.onnx")
LOCAL_TTS_TOKENS = os.path.join(MODELS_DIR, "tts_vits_tokens.txt")
LOCAL_TTS_LEXICON = os.path.join(MODELS_DIR, "tts_vits_lexicon.txt")
LLAMA_SERVER = os.path.join(RES_DIR, "dl", "llama", "llama-server.exe")
if not os.path.isfile(LLAMA_SERVER):
    LLAMA_SERVER = os.path.join(RES_DIR, "bin", "llama-server.exe")


def _find_local_llm_model():
    """自动探测本地大模型 GGUF：优先 dl/models 与 dl 下最大的主模型（排除 mmproj 多模态投影文件）。
    用户放入任意量化版本（如 Qwen3.6-35B-A3B 系列）均可直接识别，无需改代码；找不到则本地模型不可用。"""
    cands = []
    for d in (os.path.join(RES_DIR, "dl"), os.path.join(RES_DIR, "models")):
        if os.path.isdir(d):
            for f in os.listdir(d):
                low = f.lower()
                if low.endswith(".gguf") and not low.startswith("mmproj"):
                    p = os.path.join(d, f)
                    try:
                        cands.append((os.path.getsize(p), p))
                    except Exception:
                        pass
    if cands:
        cands.sort(key=lambda x: x[0], reverse=True)
        return cands[0][1]
    return ""


def _find_local_mmproj():
    """自动探测多模态投影文件 mmproj-*.gguf（Qwen3.6-35B 视觉必需）"""
    for d in (os.path.join(RES_DIR, "dl"), os.path.join(RES_DIR, "models")):
        if os.path.isdir(d):
            for f in os.listdir(d):
                if f.lower().startswith("mmproj") and f.lower().endswith(".gguf"):
                    return os.path.join(d, f)
    return None


def local_llm_model_name():
    """llama-server /v1/models 返回的模型 id = 模型文件名（不含路径）"""
    return os.path.basename(LLAMA_MODEL)


LLAMA_MODEL = _find_local_llm_model()
LLAMA_MMPROJ = _find_local_mmproj()

# ============ 配置 ============
CONFIG_PATH = os.path.join(APP_DIR, "config.json")
HISTORY_DIR = os.path.join(APP_DIR, "history")
IMAGE_DIR = os.path.join(APP_DIR, "images")
SEARCH_CACHE_DIR = os.path.join(APP_DIR, "images", "search_cache")
TTS_CACHE_DIR = os.path.join(APP_DIR, "tts_cache")
VIDEO_DIR = os.path.join(APP_DIR, "videos")

DEFAULT_CONFIG = {
    "providers": {
        "agnes": {
            "type": "openai", "name": "AgnesAI（主）",
            "base_url": "https://apihub.agnes-ai.com/v1",
            "api_key": "",
            "chat_model": "agnes-2.5-flash", "image_model": "agnes-image-2.1-flash",
            "video_model": "agnes-video-2.5",
            "enabled": True
        },
        "zhipu": {
            "type": "openai", "name": "智谱 GLM-4-Flash（免费）",
            "base_url": "https://open.bigmodel.cn/api/paas/v4",
            "api_key": "", "chat_model": "glm-4-flash", "image_model": "",
            "enabled": True
        },
        "baidu": {
            "type": "baidu", "name": "百度 ERNIE-Speed（免费）",
            "base_url": "", "api_key": "", "secret_key": "",
            "chat_model": "ernie-speed-128k", "image_model": "",
            "enabled": True
        },
        "ollama": {
            "type": "openai", "name": "本地模型（包内 Qwen3.6-35B，完全离线）",
            "base_url": "http://127.0.0.1:18080/v1",
            "api_key": "ollama", "chat_model": "Qwen3.6-35B-A3B-IQ2_M.gguf", "image_model": "",
            "enabled": True
        },
        "g4f": {
            "type": "g4f", "name": "gpt4free（GitHub 开源保底，无需Key）",
            "base_url": "", "api_key": "",
            "chat_model": "gpt-4o-mini", "image_model": "",
            "enabled": True
        },
        "freellmapi": {
            "type": "openai", "name": "FreeLLMAPI（本地免费聚合，自动调度）",
            "base_url": "http://127.0.0.1:3001/v1",
            "api_key": "", "chat_model": "auto", "image_model": "",
            "enabled": True
        },
        "github": {
            "type": "openai", "name": "GitHub Models（免费，需 GitHub PAT）",
            "base_url": "https://models.github.ai/inference",
            "api_key": "", "chat_model": "gpt-4.1-mini", "image_model": "",
            "enabled": True
        }
    },
    "mode": "cloud",
    "tts_enabled": True,
    "max_history": 20,
    "temperature": 0.3,
    "huggingface_token": "",
    "system_prompt": """你是 AgnesAgent，一款由 MIRAGE 独立开发的免费全能 AI 桌面助手。你必须严格遵循以下规则回答问题：

【核心原则】
1. 直接回答优先：先用自己的知识直接回答用户问题，不要一上来就调用工具。回答必须紧扣用户当前问题，禁止跑题、发散或自行展开无关话题。
2. 信息不足时：如果自己的知识不足以回答，先给出已有信息的部分答案，然后明确声明「关于XX我无法确认」，再考虑使用搜索工具补充。
3. 严禁编造：绝对禁止虚构、捏造或猜测答案。不确定就说不知道，不要编造数据、引用或事实。
4. 回答格式：保持简洁、结构化。能用一句话说清就不要用三段。列表、代码、步骤用对应格式呈现。
5. 工具使用：仅在用户询问实时信息（天气、新闻、股价等）、需要执行代码、或需要读取本地文件时才使用工具。不要每轮都搜索。
6. 本地能力：你可以调用本地离线 AI（无需联网）：ocr_image 识别图片文字；local_llm 调用本地大模型回答（隐私/断网场景优先）。离线优先，能用本地能力完成的任务优先用本地能力。
7. 图像生成：当用户明确要求"生成/画/绘制/做一张图/图片"（如画一只猫、生成头像、做海报图）时，必须调用 generate_image 工具生成图片，生成后简短告知用户即可，不要自行描述图片内容，也不要回复"没有此功能"。

用户可通过「📎 文件」按钮上传附件（txt/pdf/docx/xlsx/csv/md/图片等），附件内容会注入对话上下文。侧边栏「能力部门」可开关你的工具集，关闭的部门你不会使用对应功能。"""
}

# 旧配置字段 → provider 迁移
LEGACY_FIELDS = ("api_base", "api_key", "chat_model", "image_model")

# 旧版弱约束 system prompt（用于识别用户未自定义、需升级为强约束的配置）
OLD_DEFAULT_SYSTEM_PROMPT = ("你是 AgnesAgent，一个免费全能的 AI 桌面助手。你回答问题简洁准确，"
                             "会主动使用可用工具（联网搜索、执行代码、文件读取）来完成任务。")


def load_config():
    cfg = dict(DEFAULT_CONFIG)
    try:
        if os.path.exists(CONFIG_PATH):
            with open(CONFIG_PATH, "r", encoding="utf-8") as f:
                saved = json.load(f)
            cfg.update(saved)
    except Exception:
        pass
    # 升级旧版默认提示词为强约束版本（仅当用户未自定义时）
    sp = cfg.get("system_prompt", "")
    if sp == OLD_DEFAULT_SYSTEM_PROMPT or (sp and sp.startswith("你是 AgnesAgent，一个免费全能的 AI 桌面助手。你回答问题简洁准确")):
        cfg["system_prompt"] = DEFAULT_CONFIG["system_prompt"]
    # 补全新字段
    cfg.setdefault("mode", DEFAULT_CONFIG["mode"])
    cfg.setdefault("temperature", DEFAULT_CONFIG["temperature"])
    cfg.setdefault("max_history", DEFAULT_CONFIG["max_history"])
    cfg.setdefault("huggingface_token", "")
    # 旧版配置迁移：只有 api_base/api_key 没有 providers 时，合并为 agnes provider
    if "providers" not in cfg or not cfg.get("providers"):
        old = {k: cfg.get(k) for k in LEGACY_FIELDS if k in cfg}
        if old:
            cfg["providers"] = {
                "agnes": {
                    "type": "openai", "name": "AgnesAI（主）",
                    "base_url": old.get("api_base", DEFAULT_CONFIG["providers"]["agnes"]["base_url"]),
                    "api_key": old.get("api_key", DEFAULT_CONFIG["providers"]["agnes"]["api_key"]),
                    "chat_model": old.get("chat_model", DEFAULT_CONFIG["providers"]["agnes"]["chat_model"]),
                    "image_model": old.get("image_model", DEFAULT_CONFIG["providers"]["agnes"]["image_model"]),
                    "enabled": True
                }
            }
    # 补全新字段
    for pid, pdef in DEFAULT_CONFIG["providers"].items():
        p = cfg["providers"].setdefault(pid, {})
        for k, v in pdef.items():
            p.setdefault(k, v)
    return cfg


def save_config(cfg):
    try:
        with open(CONFIG_PATH, "w", encoding="utf-8") as f:
            json.dump(cfg, f, ensure_ascii=False, indent=2)
    except Exception:
        pass


def ensure_dirs():
    for d in (HISTORY_DIR, IMAGE_DIR, SEARCH_CACHE_DIR, TTS_CACHE_DIR, VIDEO_DIR):
        os.makedirs(d, exist_ok=True)


# ============ 工具函数 ============
HEADERS = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/126.0 Safari/537.36"}
import re


def _clean_html(s):
    if not s:
        return ""
    return re.sub(r"\s+", " ", re.sub(r"<[^>]+>", "", s)).strip()


def web_search(query, max_results=5):
    """联网搜索：必应 → 百度 → DuckDuckGo 三层兜底，返回结构化列表"""
    results = []
    try:
        r = requests.get("https://www.bing.com/search",
                         params={"q": query, "count": str(max_results), "setlang": "zh-hans"},
                         headers=HEADERS, timeout=12)
        r.raise_for_status()
        blocks = re.findall(r'<li class="b_algo".*?</li>', r.text, re.S)
        for b in blocks[:max_results]:
            title_m = re.search(r'<h2[^>]*>\s*<a[^>]*href="([^"]+)"[^>]*>(.*?)</a>', b, re.S)
            snip_m = re.search(r'<p[^>]*>(.*?)</p>', b, re.S)
            if title_m:
                results.append({
                    "title": _clean_html(title_m.group(2)),
                    "url": title_m.group(1),
                    "snippet": _clean_html(snip_m.group(1)) if snip_m else ""
                })
    except Exception:
        pass
    # 百度兜底
    if not results:
        try:
            r2 = requests.get("https://www.baidu.com/s", params={"wd": query}, headers=HEADERS, timeout=12)
            r2.raise_for_status()
            blocks2 = re.findall(r'<div[^>]*class="[^"]*result[^"]*c-container[^"]*".*?</div>\s*</div>', r2.text, re.S)
            for b in blocks2[:max_results]:
                title_m = re.search(r'<h3[^>]*>\s*<a[^>]*href="([^"]+)"[^>]*>(.*?)</a>', b, re.S)
                if title_m:
                    url = title_m.group(1)
                    if url.startswith("http://www.baidu.com/link?url="):
                        url = "https://www.baidu.com/s?wd=" + query + "#" + url
                    snip_m = re.search(r'<span[^>]*class="[^"]*content-right[^"]*"[^>]*>(.*?)</span>', b, re.S) or \
                             re.search(r'<div[^>]*class="[^"]*c-span-last[^"]*"[^>]*>(.*?)</div>', b, re.S)
                    results.append({
                        "title": _clean_html(title_m.group(2)),
                        "url": url,
                        "snippet": _clean_html(snip_m.group(1)) if snip_m else ""
                    })
        except Exception:
            pass
    # DuckDuckGo 兜底（ddgs 开源库，免 Key，多后端轮询）
    if not results:
        try:
            if DDGS is not None:
                with DDGS() as ddgs:
                    for backend in ("duckduckgo", "brave", "bing"):
                        try:
                            for r in ddgs.text(query, max_results=max_results, backend=backend):
                                results.append({"title": r.get("title", ""),
                                                "url": r.get("href", ""),
                                                "snippet": r.get("body", "")})
                            if results:
                                break
                        except Exception:
                            continue
        except Exception:
            pass
    # DuckDuckGo HTML 手搓兜底（ddgs 库不可用时）
    if not results:
        try:
            r3 = requests.post("https://html.duckduckgo.com/html/",
                               data={"q": query}, headers=HEADERS, timeout=12)
            r3.raise_for_status()
            blocks3 = re.findall(r'<div class="result results_links.*?</div>\s*</div>', r3.text, re.S)
            for b in blocks3[:max_results]:
                title_m = re.search(r'<a[^>]*class="result__a"[^>]*href="([^"]+)"[^>]*>(.*?)</a>', b, re.S)
                snip_m = re.search(r'<a[^>]*class="result__snippet"[^>]*>(.*?)</a>', b, re.S)
                if title_m:
                    url = title_m.group(1)
                    if url.startswith("//"):
                        url = "https:" + url
                    results.append({
                        "title": _clean_html(title_m.group(2)),
                        "url": url,
                        "snippet": _clean_html(snip_m.group(1)) if snip_m else ""
                    })
        except Exception:
            pass
    return results


def web_search_images(query, max_results=9):
    """图片搜索：Bing Images 解析缩略图与原图 URL，返回结构化列表"""
    results = []
    try:
        r = requests.get("https://www.bing.com/images/search",
                         params={"q": query, "count": str(max_results), "qft": "+filterui:imagesize-medium"},
                         headers=HEADERS, timeout=15)
        r.raise_for_status()
        # 解析 iusc 块的 m 属性 JSON（含 murl 原图 / turl 缩略图）
        import html as _html
        pat = re.compile(r'class="iusc"[^>]*?m="([^"]+)"')
        for m in pat.findall(r.text):
            try:
                data = json.loads(_html.unescape(m))
                murl = data.get("murl", "")
                turl = data.get("turl", "") or data.get("cturl", "")
                if murl:
                    results.append({"title": data.get("t", "") or query, "thumb": turl or murl, "url": murl})
            except Exception:
                continue
            if len(results) >= max_results:
                break
    except Exception:
        pass
    # ddgs 图片兜底（Bing 失败时，多后端轮询）
    if not results:
        try:
            if DDGS is not None:
                with DDGS() as ddgs:
                    for backend in ("duckduckgo", "brave", "bing"):
                        try:
                            for r in ddgs.images(query, max_results=max_results, backend=backend):
                                results.append({"title": r.get("title", "") or query,
                                                "thumb": r.get("thumbnail", ""),
                                                "url": r.get("image", "")})
                            if results:
                                break
                        except Exception:
                            continue
        except Exception:
            pass
    return results


def run_python(code):
    """安全执行 Python 代码（受限沙箱），返回 stdout / 结果"""
    import ast
    out = io.StringIO()
    old_stdout = sys.stdout
    sys.stdout = out
    try:
        tree = ast.parse(code)
        banned = ["os", "sys", "subprocess", "shutil", "socket", "ctypes", "importlib", "builtins", "requests"]
        for node in ast.walk(tree):
            if isinstance(node, ast.Import):
                for alias in node.names:
                    if alias.name.split(".")[0] in banned:
                        return f"禁止导入模块: {alias.name}"
            elif isinstance(node, ast.ImportFrom):
                if (node.module or "").split(".")[0] in banned:
                    return f"禁止导入模块: {node.module}"
        safe_builtins = {"print": print, "len": len, "range": range, "int": int, "float": float,
                         "str": str, "list": list, "dict": dict, "tuple": tuple, "set": set,
                         "sum": sum, "min": min, "max": max, "abs": abs, "round": round,
                         "enumerate": enumerate, "zip": zip, "sorted": sorted, "reversed": reversed,
                         "bool": bool, "type": type, "isinstance": isinstance, "pow": pow, "divmod": divmod}
        namespace = {"__builtins__": safe_builtins}
        exec(compile(tree, "<sandbox>", "exec"), namespace)
        return out.getvalue().strip() or "(执行完成，无输出)"
    except Exception as e:
        return f"执行出错: {type(e).__name__}: {e}"
    finally:
        sys.stdout = old_stdout


def read_local_file(path):
    """读取本地文件内容：文本 / PDF / Word / Excel 多格式（基于开源库 pypdf / python-docx / openpyxl）"""
    try:
        if not os.path.exists(path):
            return f"文件不存在: {path}"
        size = os.path.getsize(path)
        if size > 10 * 1024 * 1024:
            return f"文件过大({size//1024//1024}MB)，仅支持 10MB 以内的文件"
        ext = os.path.splitext(path)[1].lower()
        if ext == ".pdf" and PdfReader is not None:
            reader = PdfReader(path)
            parts = []
            for i, page in enumerate(reader.pages[:50]):
                parts.append(f"--- 第{i+1}页 ---\n" + (page.extract_text() or ""))
            content = "\n".join(parts)
        elif ext == ".docx" and DocxDocument is not None:
            doc = DocxDocument(path)
            content = "\n".join(p.text for p in doc.paragraphs[:500])
        elif ext in (".xlsx", ".xlsm") and load_workbook is not None:
            wb = load_workbook(path, read_only=True, data_only=True)
            parts = []
            for ws in wb.worksheets[:5]:
                parts.append(f"--- Sheet: {ws.title} ---")
                for row in ws.iter_rows(max_row=200, values_only=True):
                    parts.append("\t".join("" if c is None else str(c) for c in row))
            content = "\n".join(parts)
            wb.close()
        else:
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                content = f.read()
        return content[:30000] + ("\n...(已截断)" if len(content) > 30000 else "")
    except Exception as e:
        return f"读取失败: {e}"


def list_files(path, pattern="*", recursive=False):
    """列出目录下的文件（支持通配符过滤与递归）"""
    if not os.path.exists(path):
        return f"路径不存在: {path}"
    if os.path.isfile(path):
        return f"{path} 是一个文件，请传入目录路径"
    import fnmatch
    try:
        entries = []
        if recursive:
            for root, dirs, files in os.walk(path):
                for fn in files:
                    if fnmatch.fnmatch(fn, pattern):
                        full = os.path.join(root, fn)
                        try:
                            size = os.path.getsize(full)
                        except Exception:
                            size = 0
                        entries.append(f"{full} ({size//1024}KB)")
        else:
            for fn in os.listdir(path):
                full = os.path.join(path, fn)
                if os.path.isfile(full) and fnmatch.fnmatch(fn, pattern):
                    try:
                        size = os.path.getsize(full)
                    except Exception:
                        size = 0
                    entries.append(f"{full} ({size//1024}KB)")
        if not entries:
            return f"未找到匹配 pattern={pattern} 的文件"
        return f"共 {len(entries)} 个文件：\n" + "\n".join(entries[:200])
    except Exception as e:
        return f"列目录失败: {e}"


def write_local_file(path, content):
    """将文本内容写入本地文件（自动创建目录）"""
    try:
        if len(content) > 2 * 1024 * 1024:
            return "内容过大(超过2MB)，请分次写入"
        path = os.path.abspath(path)
        d = os.path.dirname(path)
        if d and not os.path.exists(d):
            os.makedirs(d, exist_ok=True)
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
        return f"已写入: {path}（{len(content)} 字符）"
    except Exception as e:
        return f"写入失败: {e}"


def file_convert(src_path, dst_path):
    """文件格式转换：pdf→txt / docx→txt / xlsx→csv / csv→xlsx / txt→md"""
    try:
        if not os.path.exists(src_path):
            return f"源文件不存在: {src_path}"
        src_ext = os.path.splitext(src_path)[1].lower()
        dst_ext = os.path.splitext(dst_path)[1].lower()
        d = os.path.dirname(os.path.abspath(dst_path))
        if d and not os.path.exists(d):
            os.makedirs(d, exist_ok=True)
        if src_ext == ".pdf" and dst_ext == ".txt" and PdfReader is not None:
            reader = PdfReader(src_path)
            text = "\n".join((page.extract_text() or "") for page in reader.pages)
            with open(dst_path, "w", encoding="utf-8") as f:
                f.write(text)
        elif src_ext == ".docx" and dst_ext == ".txt" and DocxDocument is not None:
            doc = DocxDocument(src_path)
            text = "\n".join(p.text for p in doc.paragraphs)
            with open(dst_path, "w", encoding="utf-8") as f:
                f.write(text)
        elif src_ext == ".xlsx" and dst_ext == ".csv" and load_workbook is not None:
            import csv as _csv
            wb = load_workbook(src_path, read_only=True, data_only=True)
            ws = wb.active
            with open(dst_path, "w", encoding="utf-8-sig", newline="") as f:
                w = _csv.writer(f)
                for row in ws.iter_rows(values_only=True):
                    w.writerow(["" if c is None else c for c in row])
            wb.close()
        elif src_ext == ".csv" and dst_ext == ".xlsx" and load_workbook is not None:
            from openpyxl import Workbook
            import csv as _csv
            wb = Workbook()
            ws = wb.active
            with open(src_path, "r", encoding="utf-8-sig", errors="replace") as f:
                for row in _csv.reader(f):
                    ws.append(row)
            wb.save(dst_path)
        elif src_ext == ".txt" and dst_ext == ".md":
            with open(src_path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
            with open(dst_path, "w", encoding="utf-8") as f:
                f.write(text)
        else:
            return f"暂不支持 {src_ext} → {dst_ext} 转换（支持 pdf→txt, docx→txt, xlsx→csv, csv→xlsx, txt→md）"
        return f"转换完成: {dst_path}"
    except Exception as e:
        return f"转换失败: {e}"


# ============ 扩充能力：系统命令 / 天气 / 语音 ============

# 危险命令黑名单：任何包含以下关键字的命令一律拒绝执行（防误删/防破坏）
SHELL_BLACKLIST = [
    "format", "del /", "rm -", "rmdir", "rd /", "deltree", "reg delete", "reg add",
    "taskkill", "shutdown", "restart-computer", "stop-computer", "remove-item",
    "clear-recyclebin", "net user", "net localgroup", "diskpart", "bcdedit",
    ">nul", "format c:", "mkfs", "dd if=", ":(){", "wmic process delete",
]

def run_shell(command, timeout=30):
    """执行系统 Shell 命令（受限）：返回命令输出。
    危险命令（删除/格式化/关机/注册表/进程结束等）会被拒绝。"""
    if not command or not command.strip():
        return "命令为空"
    low = command.lower()
    for kw in SHELL_BLACKLIST:
        if kw in low:
            return f"拒绝执行：命令包含危险操作关键词 [{kw}]，已拦截。仅允许安全的只读/查询/创建类命令。"
    try:
        r = subprocess.run(command, shell=True, capture_output=True, text=True,
                           timeout=timeout, encoding="utf-8", errors="replace")
        out = (r.stdout or "").strip()
        err = (r.stderr or "").strip()
        parts = []
        if out:
            parts.append(out)
        if err:
            parts.append(f"[stderr] {err}")
        if not parts:
            return f"命令执行完成，退出码 {r.returncode}，无输出"
        if r.returncode != 0:
            return "\n".join(parts) + f"\n[退出码 {r.returncode}]"
        return "\n".join(parts)[:8000]
    except subprocess.TimeoutExpired:
        return f"命令执行超时（>{timeout}s），已终止"
    except Exception as e:
        return f"命令执行异常: {e}"


def get_weather(city="北京"):
    """免费天气查询（wttr.in，无需 API Key）"""
    try:
        url = f"https://wttr.in/{requests.utils.quote(city)}?format=j1&lang=zh"
        r = requests.get(url, timeout=20, headers={"User-Agent": "curl/8.0"})
        if r.status_code != 200:
            return f"天气查询失败（HTTP {r.status_code}）"
        data = r.json()
        cur = data.get("current_condition", [{}])[0]
        weather = cur.get("lang_zh", [{}])
        desc = weather[0].get("value", cur.get("weatherDesc", [{}])[0].get("value", "未知")) if weather else cur.get("weatherDesc", [{}])[0].get("value", "未知")
        temp = cur.get("temp_C", "?")
        feels = cur.get("FeelsLikeC", "?")
        hum = cur.get("humidity", "?")
        wind = cur.get("windspeedKmph", "?")
        city_name = data.get("nearest_area", [{}])[0].get("areaName", [{}])[0].get("value", city)
        return (f"{city_name} 当前天气：{desc}，气温 {temp}°C（体感 {feels}°C），"
                f"湿度 {hum}%，风速 {wind}km/h")
    except Exception as e:
        return f"天气查询失败: {e}"


# ============ 语音输入（录音 + 免费 ASR）============
_SOUNDDEVICE_OK = False
_SPEECHREC_OK = False
try:
    import sounddevice as sd
    _SOUNDDEVICE_OK = True
except Exception:
    sd = None
try:
    import speech_recognition as sr
    _SPEECHREC_OK = True
except Exception:
    sr = None

def record_audio_to_wav(path, duration=6, sample_rate=16000):
    """用默认麦克风录音，保存为 WAV（16kHz 单声道）。返回 (成功?, 说明)"""
    if not _SOUNDDEVICE_OK:
        return False, "未安装 sounddevice，请运行 pip install sounddevice"
    try:
        import numpy as np
        frames = sd.rec(int(duration * sample_rate), samplerate=sample_rate, channels=1, dtype="int16")
        sd.wait()
        data = frames.reshape(-1)
        with wave_open(path, "wb") as wf:
            wf.setnchannels(1)
            wf.setsampwidth(2)
            wf.setframerate(sample_rate)
            wf.writeframes(data.tobytes())
        return True, path
    except Exception as e:
        return False, f"录音失败: {e}"

def wave_open(path, mode):
    import wave
    return wave.open(path, mode)

def transcribe_audio(path):
    """免费语音转文字：优先 Google Web Speech（免 Key）。失败降级：sherpa-onnx 本地离线 → Sphinx 离线。"""
    # 优先本地离线 ASR（sherpa-onnx SenseVoice，效果好且不依赖网络）
    try:
        local_text, local_err = local_asr(path)
        if local_text:
            return local_text, ""
    except Exception:
        pass
    if not _SPEECHREC_OK:
        return "", "未安装 SpeechRecognition，请运行 pip install SpeechRecognition"
    r = sr.Recognizer()
    try:
        with sr.AudioFile(path) as src:
            audio = r.record(src)
    except Exception as e:
        return "", f"音频读取失败: {e}"
    # 优先 Google（免费、支持中文、效果最好）
    for engine in ("google", "sphinx"):
        try:
            if engine == "google":
                return r.recognize_google(audio, language="zh-CN"), ""
            return r.recognize_sphinx(audio, language="zh-CN"), ""
        except sr.UnknownValueError:
            if engine == "google":
                continue
            return "", "未能识别出语音内容（可尝试更清晰的普通话）"
        except sr.RequestError as e:
            if engine == "google":
                continue
            return "", f"语音识别服务不可用: {e}"
    return "", "语音识别失败"


# ============ 本地 AI 能力（离线开源算法，全部防御式）============
# ---- 1. 本地 LLM（llama-server 子进程 + OpenAI 兼容 API）----
_LLAMA_PROCESS = None
_LLAMA_BASE_URL = "http://127.0.0.1:18080/v1"


def _gpu_offload_layers():
    """按 NVIDIA 显存大小保守估算可 offload 的层数（无独显/探测失败返回 0 走纯 CPU）"""
    try:
        out = subprocess.run(
            ["nvidia-smi", "--query-gpu=memory.total", "--format=csv,noheader,nounits"],
            capture_output=True, text=True, timeout=5)
        mb = int(out.stdout.strip().splitlines()[0].strip())
        if mb >= 20000:
            return 99
        if mb >= 12000:
            return 60
        if mb >= 8000:
            return 40
        if mb >= 4000:
            return 24
        if mb >= 2000:
            return 10
    except Exception:
        pass
    return 0


def _start_llama_server():
    """启动 llama-server 子进程（仅一次，lazy 启动）"""
    global _LLAMA_PROCESS
    if _LLAMA_PROCESS is not None:
        return True
    if not os.path.isfile(LLAMA_SERVER) or not os.path.isfile(LLAMA_MODEL):
        return False
    # 先探测端口：已有可用本地服务则直接复用，避免重复拉起与冷启动超时静默降级
    try:
        r = requests.get(f"{_LLAMA_BASE_URL}/models", timeout=2)
        if r.status_code == 200:
            _LLAMA_PROCESS = "reuse"
            return True
    except Exception:
        pass
    try:
        ngl = _gpu_offload_layers()
        cmd = [LLAMA_SERVER, "-m", LLAMA_MODEL, "--host", "127.0.0.1", "--port", "18080",
               "-c", "8192", "-ngl", str(ngl), "--no-mmap"]
        if LLAMA_MMPROJ:
            cmd += ["--mmproj", LLAMA_MMPROJ]
        cmd += ["--jinja"]  # Qwen 新模板必需，不加会回复异常/中文乱码
        _LLAMA_PROCESS = subprocess.Popen(
            cmd,
            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
            creationflags=subprocess.CREATE_NO_WINDOW
        )
        # 等待服务就绪（最多 180 秒，大模型冷加载较慢）
        for _ in range(180):
            try:
                r = requests.get(f"{_LLAMA_BASE_URL}/models", timeout=3)
                if r.status_code == 200:
                    return True
            except Exception:
                pass
            time.sleep(1)
        return False
    except Exception:
        return False

def local_llm_chat(messages, max_tokens=1024, temperature=0.3):
    """调用本地 llama-server 进行对话（离线，完全免费，无需任何网络）"""
    if not _start_llama_server():
        return "本地 LLM 服务未就绪（模型文件或 llama-server 缺失）"
    try:
        r = requests.post(f"{_LLAMA_BASE_URL}/chat/completions", json={
            "model": local_llm_model_name(), "messages": messages,
            "max_tokens": max_tokens, "temperature": temperature,
            "stream": False,
            "chat_template_kwargs": {"enable_thinking": False},  # 关闭思考模式，直接输出正文，避免 UI 空白
        }, timeout=300)
        r.raise_for_status()
        data = r.json()
        return data["choices"][0]["message"]["content"]
    except Exception as e:
        return f"本地 LLM 出错: {e}"

def local_llm_stream(messages):
    """流式调用本地 llama-server"""
    if not _start_llama_server():
        yield "本地 LLM 服务未就绪"
        return
    try:
        r = requests.post(f"{_LLAMA_BASE_URL}/chat/completions", json={
            "model": local_llm_model_name(), "messages": messages,
            "temperature": 0.3, "stream": True,
            "chat_template_kwargs": {"enable_thinking": False},  # 关闭思考模式，直接输出正文，避免 UI 空白
        }, stream=True, timeout=300)
        r.raise_for_status()
        for line in r.iter_lines():
            if line:
                line = line.decode("utf-8", errors="replace").strip()
                if line.startswith("data: "):
                    data = line[6:]
                    if data == "[DONE]":
                        break
                    try:
                        chunk = json.loads(data)
                        delta = chunk.get("choices", [{}])[0].get("delta", {}).get("content", "")
                        if delta:
                            yield delta
                    except Exception:
                        pass
    except Exception as e:
        yield f"流式调用失败: {e}"

# ---- 2. 本地 OCR（RapidOCR，离线图片文字识别）----
_OCR_INSTANCE = None

def ocr_image(path):
    """用 RapidOCR 离线识别图片中的文字，返回 (行列表, 错误)"""
    global _OCR_INSTANCE
    if not RAPIDOCR_AVAILABLE:
        return [], "RapidOCR 未安装，请运行 pip install rapidocr-onnxruntime"
    if not os.path.isfile(path):
        return [], f"图片文件不存在: {path}"
    try:
        if _OCR_INSTANCE is None:
            _OCR_INSTANCE = _RapidOCR()
        result, elapse = _OCR_INSTANCE(path)
        if result is None:
            return [], "未识别到文字"
        lines = []
        for box, text, score in result:
            if text and text.strip():
                lines.append({"text": text.strip(), "confidence": round(float(score), 3)})
        return lines, ""
    except Exception as e:
        return [], f"OCR 识别失败: {e}"

# ---- 3. 本地 ASR（sherpa-onnx SenseVoice，离线语音转文字）----
_ASR_RECOGNIZER = None

def local_asr(audio_path):
    """用 sherpa-onnx SenseVoice 离线识别语音（替代 Google Web Speech 兜底）"""
    global _ASR_RECOGNIZER
    if not SHERPA_ONNX_AVAILABLE:
        return "", "sherpa-onnx 未安装，请运行 pip install sherpa-onnx"
    if not os.path.isfile(LOCAL_ASR_MODEL):
        return "", f"ASR 模型文件不存在: {LOCAL_ASR_MODEL}"
    if not os.path.isfile(audio_path):
        return "", f"音频文件不存在: {audio_path}"
    try:
        if _ASR_RECOGNIZER is None:
            _ASR_RECOGNIZER = sherpa_onnx.OfflineRecognizer.from_sense_voice(
                model=LOCAL_ASR_MODEL,
                tokens=LOCAL_ASR_TOKENS,
                use_itn=True,
                debug=False,
            )
        import soundfile as sf
        audio, sample_rate = sf.read(audio_path, dtype="float32", always_2d=True)
        audio = audio[:, 0]
        stream = _ASR_RECOGNIZER.create_stream()
        stream.accept_waveform(sample_rate, audio)
        _ASR_RECOGNIZER.decode_stream(stream)
        text = stream.result.text.strip()
        if text:
            return text, ""
        return "", "未识别出语音内容"
    except Exception as e:
        return "", f"本地 ASR 识别失败: {e}"

# ---- 4. 本地 TTS（sherpa-onnx vits-zh，离线语音合成）----
def local_tts(text, output_path):
    """用 sherpa-onnx vits-zh 离线合成中文语音，保存为 WAV 文件"""
    if not SHERPA_ONNX_AVAILABLE:
        return False, "sherpa-onnx 未安装"
    if not os.path.isfile(LOCAL_TTS_MODEL):
        return False, f"TTS 模型文件不存在: {LOCAL_TTS_MODEL}"
    try:
        tts_config = sherpa_onnx.OfflineTtsConfig(
            model=sherpa_onnx.OfflineTtsModelConfig(
                vits=sherpa_onnx.OfflineTtsVitsModelConfig(
                    model=LOCAL_TTS_MODEL,
                    tokens=LOCAL_TTS_TOKENS,
                    lexicon=LOCAL_TTS_LEXICON if os.path.isfile(LOCAL_TTS_LEXICON) else "",
                    noise_scale=0.667,
                    noise_scale_w=0.8,
                    length_scale=1.0
                ),
                num_threads=2,
                debug=False
            )
        )
        tts = sherpa_onnx.OfflineTts(tts_config)
        audio = tts.generate(text, sid=0, speed=1.0)
        if audio is not None and len(audio.samples) > 0:
            import wave
            with wave.open(output_path, "wb") as wf:
                wf.setnchannels(1)
                wf.setsampwidth(2)
                wf.setframerate(audio.sample_rate)
                import numpy as np
                wf.writeframes(np.array(audio.samples, dtype=np.int16).tobytes())
            return True, output_path
        return False, "TTS 生成失败：无音频输出"
    except Exception as e:
        return False, f"本地 TTS 失败: {e}"


TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "web_search",
            "description": "联网搜索互联网文字信息，返回搜索结果（标题、链接、摘要）。仅在用户询问实时/最新信息（新闻、天气、股价、事件等）或你确认自己的知识不足时使用；能用知识直接回答的常识问题不要调用本工具。",
            "parameters": {
                "type": "object",
                "properties": {"query": {"type": "string", "description": "搜索关键词"}},
                "required": ["query"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "web_search_images",
            "description": "联网搜索图片，返回图片链接列表。用户想找图片、看照片、查图时使用。",
            "parameters": {
                "type": "object",
                "properties": {"query": {"type": "string", "description": "图片搜索关键词"}},
                "required": ["query"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "run_python",
            "description": "执行 Python 代码做计算、数据处理、文件分析。代码中可用 print 输出结果。禁止导入 os/sys/subprocess 等系统模块。",
            "parameters": {
                "type": "object",
                "properties": {"code": {"type": "string", "description": "要执行的 Python 代码"}},
                "required": ["code"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "read_local_file",
            "description": "读取本地文件内容（需提供绝对路径）。支持文本(txt/md/py/json/csv等)、PDF、Word(.docx)、Excel(.xlsx) 多格式，用于分析用户电脑上的文档。",
            "parameters": {
                "type": "object",
                "properties": {"path": {"type": "string", "description": "文件绝对路径"}},
                "required": ["path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "list_files",
            "description": "列出目录下的文件（需提供目录绝对路径）。支持通配符过滤（如 *.pdf、*报告*）和递归子目录，用于查找、浏览用户电脑上的文件。",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "目录绝对路径"},
                    "pattern": {"type": "string", "description": "文件名通配符过滤，默认 * 全部"},
                    "recursive": {"type": "boolean", "description": "是否递归子目录，默认 false"}
                },
                "required": ["path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "write_local_file",
            "description": "将文本内容写入本地文件（需提供绝对路径，自动创建目录）。用于生成报告、脚本、笔记等文件。",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "目标文件绝对路径"},
                    "content": {"type": "string", "description": "要写入的文本内容"}
                },
                "required": ["path", "content"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "file_convert",
            "description": "文件格式转换（需提供源文件与目标文件绝对路径）。支持 pdf→txt、docx→txt、xlsx→csv、csv→xlsx、txt→md。",
            "parameters": {
                "type": "object",
                "properties": {
                    "src_path": {"type": "string", "description": "源文件绝对路径"},
                    "dst_path": {"type": "string", "description": "目标文件绝对路径（扩展名决定目标格式）"}
                },
                "required": ["src_path", "dst_path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "run_shell",
            "description": "执行系统 Shell 命令（Windows PowerShell）。只读/查询/创建类命令可用（如查看磁盘、目录列表、获取系统信息）。删除/格式化/关机/注册表修改/进程结束等危险命令已被内置黑名单拦截。",
            "parameters": {
                "type": "object",
                "properties": {"command": {"type": "string", "description": "要执行的命令"}},
                "required": ["command"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "get_weather",
            "description": "查询指定城市当前天气（气温/体感/湿度/风速），免费无需 API Key。",
            "parameters": {
                "type": "object",
                "properties": {"city": {"type": "string", "description": "城市名，如 北京/上海/广州"}},
                "required": ["city"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "ocr_image",
            "description": "离线识别图片中的文字（RapidOCR 开源引擎，无需联网）。用户想提取图片/截图/扫描件里的文字时使用。",
            "parameters": {
                "type": "object",
                "properties": {"path": {"type": "string", "description": "图片文件绝对路径（png/jpg/bmp）"}},
                "required": ["path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "local_llm",
            "description": "调用本地大模型（Qwen3.6-35B，完全离线）回答、写作或总结。用于断网/离线场景，或用户明确要求本地处理数据隐私敏感内容时。",
            "parameters": {
                "type": "object",
                "properties": {"prompt": {"type": "string", "description": "发给本地模型的提示词（可含上下文）"}},
                "required": ["prompt"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "generate_image",
            "description": "根据文字描述生成一张图片（文生图）。自动选择通道：优先 AgnesAI 免费图像 API，失败时降级免 Key 的 Pollinations.ai，再失败降级本地 sd.cpp 离线生成。用户想要图片、插画、头像、海报、示意图时使用。",
            "parameters": {
                "type": "object",
                "properties": {"prompt": {"type": "string", "description": "图片内容描述（中文，描述主体、风格、色彩、构图等细节）"}},
                "required": ["prompt"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "generate_video",
            "description": "根据文字描述生成视频（文生视频）。自动选择通道：优先 AgnesAI 免费视频 API，额度不足或失败时自动降级为本地 Wan2.1 开源模型生成（纯 CPU，较慢约 20-40 分钟）。用户想要视频、动画、短片时使用。",
            "parameters": {
                "type": "object",
                "properties": {
                    "prompt": {"type": "string", "description": "视频内容描述（中文，描述主体、动作、场景、氛围）"},
                    "seconds": {"type": "integer", "description": "视频时长秒数，默认 5，可选 5"}
                },
                "required": ["prompt"]
            }
        }
    }
]


# ============ 部门（模块化能力挂载）============
# 每个"部门" = 一个独立功能挂载点，UI 侧边栏可独立开关；
# "tools" 对应大模型可调用的工具名（关闭则不传给大模型）；
# "actions" 对应 UI 快捷按钮动作 id（关闭则按钮禁用），实现"每个功能一个部门"。
DEPARTMENTS = {
    "读取文件部": {"icon": "📂", "desc": "读取本地文件内容（txt/pdf/docx/xlsx/csv/md/代码）", "tools": ["read_local_file"], "actions": []},
    "查找文件部": {"icon": "🔎", "desc": "列出目录、查找本地文件", "tools": ["list_files"], "actions": []},
    "写入文件部": {"icon": "✍️", "desc": "创建/写入本地文件", "tools": ["write_local_file"], "actions": []},
    "转换文件部": {"icon": "🔄", "desc": "本地文件格式转换（PDF/Word/Excel/图片等）", "tools": ["file_convert"], "actions": []},
    "网页搜索部": {"icon": "🌐", "desc": "联网搜索网页资讯（必应/百度/DDG 三层兜底）", "tools": ["web_search"], "actions": ["search"]},
    "图片搜索部": {"icon": "🖼", "desc": "联网搜索图片", "tools": ["web_search_images"], "actions": ["img_search"]},
    "天气查询部": {"icon": "⛅", "desc": "查询天气", "tools": ["get_weather"], "actions": []},
    "代码执行部": {"icon": "⚙️", "desc": "执行 Python 代码（计算/数据处理/文件分析）", "tools": ["run_python"], "actions": []},
    "系统命令部": {"icon": "🖥️", "desc": "执行系统 Shell 命令（受限安全黑名单）", "tools": ["run_shell"], "actions": []},
    "OCR识别部": {"icon": "📄", "desc": "离线 OCR 图片文字识别（RapidOCR，无需联网）", "tools": ["ocr_image"], "actions": []},
    "本地大模型部": {"icon": "🧠", "desc": "本地离线大模型 Qwen3.6-35B（隐私/断网场景优先）", "tools": ["local_llm"], "actions": []},
    "云端生图部": {"icon": "🎨", "desc": "AgnesAI 免费图像 API（需联网，约 500 次/月）", "tools": ["generate_image"], "actions": ["gen_image_cloud"]},
    "本地生图部": {"icon": "🖌️", "desc": "sd.cpp + Realistic Vision V6（纯 CPU 离线，约 1-2 分钟/张）", "tools": [], "actions": ["gen_image_local"]},
    "云端生视频部": {"icon": "🎬", "desc": "AgnesAI 免费视频 API（需联网，失败自动切本地 Wan2.1）", "tools": ["generate_video"], "actions": ["gen_video_cloud"]},
    "本地生视频部": {"icon": "🎞️", "desc": "ComfyUI + Wan2.1 GGUF（本地 4.1GB 模型，纯 CPU 自动调度，约 20-40 分钟/段）", "tools": [], "actions": ["gen_video_local"]},
    "模型市场部": {"icon": "🛒", "desc": "模型市场：一键下载/更新 Wan2.1 系列模型到本地 ComfyUI", "tools": [], "actions": ["model_market"]},
    "视觉分析部": {"icon": "👁️", "desc": "上传图片进行理解问答", "tools": [], "actions": ["analyze_image"]},
    "语音输入部": {"icon": "🎤", "desc": "麦克风录音 + 免费 ASR 转文字", "tools": [], "actions": ["voice_input"]},
    "语音播报部": {"icon": "🔊", "desc": "AI 回复自动 TTS 语音播报", "tools": [], "actions": []},
}


def get_enabled_actions(dept_state=None):
    """根据部门开关状态返回可用的 UI 动作 id 集合（全部开启时返回全部）"""
    all_actions = [a for info in DEPARTMENTS.values() for a in info.get("actions", [])]
    if not dept_state or all(v for v in dept_state.values()):
        return set(all_actions)
    enabled = set()
    for dept, info in DEPARTMENTS.items():
        if dept_state.get(dept, True):
            enabled.update(info.get("actions", []))
    return enabled


def get_enabled_tools(dept_state=None):
    """根据部门开关状态返回启用的工具列表（全部开启时返回完整 TOOLS）"""
    if not dept_state or all(v for v in dept_state.values()):
        return TOOLS
    disabled = set()
    for dept, state in dept_state.items():
        if not state and dept in DEPARTMENTS:
            disabled.update(DEPARTMENTS[dept]["tools"])
    if not disabled:
        return TOOLS
    return [t for t in TOOLS if t["function"]["name"] not in disabled]


# 动作 id → (按钮文案, 处理函数)，对应 DEPARTMENTS 中各部门的 actions
ACTION_MAP = {
    "search": ("🔍 网页搜索", lambda w: w.quick_action("帮我搜索：")),
    "img_search": ("🖼 搜图片", lambda w: w.quick_action("帮我搜图片：")),
    "gen_image_cloud": ("🎨 云端生图", lambda w: w.ask_image_prompt()),
    "gen_image_local": ("🖌️ 本地生图", lambda w: w.ask_image_prompt_local()),
    "gen_video_cloud": ("🎬 云端生视频", lambda w: w.ask_video_prompt()),
    "gen_video_local": ("🎞️ 本地生视频", lambda w: w.ask_video_prompt_local()),
    "model_market": ("🛒 模型市场", lambda w: w.open_model_market()),
    "analyze_image": ("👁️ 分析图片", lambda w: w.analyze_image()),
    "voice_input": ("🎤 语音输入", lambda w: w.toggle_voice()),
}


# ============ 文生视频（AgnesAI API → 本地 Wan2.1 兜底）============
def _find_video_provider():
    """取第一个配置了 video_model 的 provider"""
    try:
        cfg = load_config()
        for p in cfg.get("providers", {}).values():
            if p.get("video_model"):
                return p
    except Exception:
        pass
    return None


def _generate_video_api(prompt, seconds=5):
    """AgnesAI 免费视频 API：提交→轮询→下载。返回 (path, err)"""
    try:
        p = _find_video_provider()
        if not p:
            return None, "未配置视频模型"
        base = p["base_url"].rstrip("/")
        key = p.get("api_key", "")
        headers = {"Authorization": f"Bearer {key}", "Content-Type": "application/json"}
        r = requests.post(base + "/video/generations",
                          json={"model": p["video_model"], "prompt": prompt, "mode": "text"},
                          headers=headers, timeout=60)
        if r.status_code != 200:
            return None, f"API 错误 {r.status_code}: {r.text[:200]}"
        task_id = r.json().get("task_id")
        if not task_id:
            return None, "API 未返回 task_id"
        deadline = time.time() + 900  # 最长等待 15 分钟
        while time.time() < deadline:
            time.sleep(10)
            try:
                rr = requests.get(base + f"/video/generations/{task_id}", headers=headers, timeout=30)
                if rr.status_code != 200:
                    continue
                d = rr.json().get("data") or rr.json()
                st = d.get("status")
                if st == "SUCCESS":
                    url = d.get("result_url")
                    if not url:
                        return None, "任务完成但无下载地址"
                    out = os.path.join(VIDEO_DIR, f"video_api_{int(time.time())}.mp4")
                    dl = requests.get(url, headers=headers, timeout=60)
                    if dl.status_code == 200 and len(dl.content) > 1000:
                        with open(out, "wb") as f:
                            f.write(dl.content)
                        return out, ""
                    return None, f"视频内容下载失败 {dl.status_code}"
                if st in ("FAILED", "ERROR", "CANCELLED"):
                    return None, f"视频任务失败: {d.get('fail_reason') or st}"
            except requests.RequestException:
                continue
        return None, "视频任务超时（15 分钟）"
    except Exception as e:
        return None, str(e)


def _generate_video_hf(prompt, seconds=5):
    """HuggingFace 免费推理 API（需注册免费 token，config.huggingface_token 填写后启用）。
    使用 Wan2.1-T2V-1.3B serverless 端点，免 GPU、免本地部署。返回 (path, err)"""
    try:
        cfg = load_config()
        token = (cfg.get("huggingface_token") or "").strip()
        if not token:
            return None, "未配置 HuggingFace token（免费注册 hf.co 后填入 config 的 huggingface_token 即启用）"
        os.makedirs(VIDEO_DIR, exist_ok=True)
        out = os.path.join(VIDEO_DIR, f"video_hf_{int(time.time())}.mp4")
        headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
        # Wan2.1-T2V-1.3B 为 HF 免费 serverless 端点；模型加载中返回 503 需重试
        last_err = ""
        for attempt in range(3):
            try:
                r = requests.post(
                    "https://api-inference.huggingface.co/models/Wan-AI/Wan2.1-T2V-1.3B",
                    json={"inputs": prompt}, headers=headers, timeout=120)
                if r.status_code == 503:
                    try:
                        d = r.json()
                        last_err = f"模型加载中（约 {d.get('estimated_time', 20)} 秒），第 {attempt + 1} 次重试"
                    except Exception:
                        last_err = f"模型加载中，第 {attempt + 1} 次重试"
                    time.sleep(15)
                    continue
                if r.status_code == 401:
                    return None, "HuggingFace token 无效或已过期"
                if r.status_code != 200:
                    last_err = f"HF 错误 {r.status_code}: {r.text[:200]}"
                    time.sleep(10)
                    continue
                if len(r.content) > 1000:
                    with open(out, "wb") as f:
                        f.write(r.content)
                    return out, "视频已生成（HuggingFace 免费推理 Wan2.1-T2V-1.3B）"
                last_err = "HF 返回内容为空"
            except requests.RequestException as e:
                last_err = f"HF 请求异常: {e}"
                time.sleep(10)
        return None, last_err or "HuggingFace 视频生成失败"
    except Exception as e:
        return None, str(e)


def _generate_video_local(prompt, seconds=5):
    """本地 Wan2.1 兜底：自动定位/启动 ComfyUI 并生成视频。返回 (path, msg)"""
    length = 81 if seconds <= 6 else 121  # Wan2.1 标准帧数：5s≈81帧，7.5s≈121帧
    return _generate_video_local_blocking(prompt, progress_cb=None, length=length)


def generate_video_blocking(prompt, seconds=5):
    """文生视频总入口：AgnesAI API 优先 → HuggingFace 免费推理 → 本地 ComfyUI + Wan2.1 兜底。返回 (path, msg)"""
    if not prompt or not prompt.strip():
        return None, "视频描述不能为空"
    path, err = _generate_video_api(prompt, seconds)
    if path:
        return path, "视频已生成（AgnesAI 免费 API）"
    path, hf_err = _generate_video_hf(prompt, seconds)
    if path:
        return path, f"AgnesAI 失败（{err}），已自动切换 HuggingFace 免费推理生成"
    path, msg = _generate_video_local(prompt, seconds)
    if path:
        return path, f"云端失败（{err}；{hf_err}）已自动切本地 Wan2.1 生成：{msg}"
    return None, f"AgnesAI：{err}；HuggingFace：{hf_err}；本地也失败：{msg}"


def generate_image_blocking(prompt):
    """文生图总入口：AgnesAI 免费 API 优先 → Pollinations 免 Key 兜底 → 本地 sd.cpp。返回 (path, msg)"""
    if not prompt or not prompt.strip():
        return None, "图片描述不能为空"
    os.makedirs(IMAGE_DIR, exist_ok=True)
    path = os.path.join(IMAGE_DIR, f"image_{time.strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:6]}.png")
    # 1. 云端 API（配置了 image_model 的 provider）
    try:
        clients = build_clients(load_config())
        client = pick_image_client(clients)
        if client is not None:
            img_bytes = client.generate_image(prompt)
            if img_bytes and len(img_bytes) > 500:
                with open(path, "wb") as f:
                    f.write(img_bytes)
                return path, "图片已生成（云端免费 API）"
    except Exception as e:
        cloud_err = str(e)
    else:
        cloud_err = "未配置支持图像生成的 provider"
    # 2. Pollinations 免 Key 兜底
    try:
        pollinations_image(prompt, path)
        if os.path.exists(path) and os.path.getsize(path) > 500:
            return path, f"云端失败（{cloud_err}），已自动切换 Pollinations 免费通道"
    except Exception as e:
        poll_err = str(e)
    else:
        poll_err = "Pollinations 未返回有效图片"
    # 3. 本地 sd.cpp 离线兜底
    try:
        generate_image_local_sd(prompt, path)
        return path, f"云端与 Pollinations 均失败，已自动切换本地 sd.cpp 离线生成（{cloud_err} | {poll_err}）"
    except Exception as e:
        return None, f"图片生成失败：云端（{cloud_err}）→ Pollinations（{poll_err}）→ 本地（{e}）"


def dispatch_tool(name, args):
    try:
        if name == "web_search":
            res = web_search(args.get("query", ""))
            return {"kind": "search", "query": args.get("query", ""), "results": res}
        elif name == "web_search_images":
            res = web_search_images(args.get("query", ""))
            return {"kind": "images", "query": args.get("query", ""), "results": res}
        elif name == "generate_image":
            path, msg = generate_image_blocking(args.get("prompt", ""))
            if path:
                return {"kind": "image", "path": path, "msg": msg}
            return {"kind": "text", "text": msg}
        elif name == "generate_video":
            path, msg = generate_video_blocking(args.get("prompt", ""), int(args.get("seconds", 5) or 5))
            if path:
                return {"kind": "video", "path": path, "msg": msg}
            return {"kind": "text", "text": msg}
        elif name == "run_python":
            return {"kind": "text", "text": run_python(args.get("code", ""))}
        elif name == "read_local_file":
            return {"kind": "text", "text": read_local_file(args.get("path", ""))}
        elif name == "list_files":
            return {"kind": "text", "text": list_files(args.get("path", ""),
                                                       args.get("pattern", "*"),
                                                       bool(args.get("recursive", False)))}
        elif name == "write_local_file":
            return {"kind": "text", "text": write_local_file(args.get("path", ""), args.get("content", ""))}
        elif name == "file_convert":
            return {"kind": "text", "text": file_convert(args.get("src_path", ""), args.get("dst_path", ""))}
        elif name == "run_shell":
            return {"kind": "text", "text": run_shell(args.get("command", ""))}
        elif name == "get_weather":
            return {"kind": "text", "text": get_weather(args.get("city", "北京"))}
        elif name == "ocr_image":
            lines, err = ocr_image(args.get("path", ""))
            if err:
                return {"kind": "text", "text": err}
            return {"kind": "text", "text": "\n".join(f"{l['text']} (置信度{l['confidence']})" for l in lines)}
        elif name == "local_llm":
            messages = [{"role": "user", "content": args.get("prompt", "")}]
            return {"kind": "text", "text": local_llm_chat(messages)}
        return {"kind": "text", "text": f"未知工具: {name}"}
    except Exception as e:
        return {"kind": "text", "text": f"工具执行异常: {e}"}


# ============ 多 Provider API 封装 ============
class APIError(Exception):
    pass


class ProviderClient:
    """单个 Provider 客户端（OpenAI 兼容 / 百度千帆 两种类型）"""

    def __init__(self, pid, cfg):
        self.pid = pid
        self.cfg = cfg
        self.type = cfg.get("type", "openai")
        self.name = cfg.get("name", pid)

    # ---------- 对话 ----------
    def chat_stream(self, messages, tools=None, tool_choice="auto"):
        if self.type == "baidu":
            yield from self._chat_baidu(messages)
            return
        if self.type == "g4f":
            yield from self._chat_g4f(messages)
            return
        url = self.cfg["base_url"].rstrip("/") + "/chat/completions"
        payload = {"model": self.cfg["chat_model"], "messages": messages,
                   "stream": True, "temperature": float(self.cfg.get("temperature", 0.3))}
        # 推理模型（agnes 等）默认开启思考，正文全在 reasoning_content 里，导致 UI 空白；
        # 实测 agnes 云端只认 chat_template_kwargs.enable_thinking（顶层 enable_thinking 无效），
        # 显式关闭思考模式，直接输出正文（服务端不识别该参数时会忽略，无副作用）
        if self.pid == "agnes":
            payload["chat_template_kwargs"] = {"enable_thinking": False}
        # 本地 llama-server（ollama provider）：Qwen3.6 等思考模型默认输出 reasoning_content、
        # content 为空，UI 只读 content 会一直空白；必须通过 chat_template_kwargs 显式关闭思考
        if self.pid == "ollama":
            payload["chat_template_kwargs"] = {"enable_thinking": False}
        if tools and self.type == "openai":
            payload["tools"] = tools
            payload["tool_choice"] = tool_choice
        headers = {"Authorization": f"Bearer {self.cfg['api_key']}",
                   "Content-Type": "application/json"}
        try:
            resp = requests.post(url, json=payload, headers=headers, stream=True, timeout=120)
        except requests.RequestException as e:
            raise APIError(f"网络错误: {e}")
        if resp.status_code != 200:
            body = resp.content.decode("utf-8", errors="replace")[:300]
            raise APIError(f"API 错误 {resp.status_code}: {body}")
        content_parts, tool_calls, current_idx = [], {}, None
        for raw_line in resp.iter_lines(decode_unicode=False):
            if not raw_line:
                continue
            line = raw_line.decode("utf-8", errors="replace")
            if not line.startswith("data:"):
                continue
            data = line[5:].strip()
            if data == "[DONE]":
                break
            try:
                chunk = json.loads(data)
            except Exception:
                continue
            if "choices" not in chunk or not chunk["choices"]:
                continue
            delta = chunk["choices"][0].get("delta", {})
            if delta.get("content"):
                content_parts.append(delta["content"])
                yield ("content", delta["content"])
            if delta.get("tool_calls"):
                for tc in delta["tool_calls"]:
                    idx = tc.get("index", 0)
                    fn = tc.get("function", {})
                    if idx not in tool_calls:
                        tool_calls[idx] = {"id": tc.get("id", f"call_{idx}"),
                                           "name": fn.get("name", ""), "arguments": ""}
                    if fn.get("name"):
                        tool_calls[idx]["name"] = fn["name"]
                    if fn.get("arguments"):
                        tool_calls[idx]["arguments"] += fn["arguments"]
        final_calls = []
        for idx in sorted(tool_calls.keys()):
            tc = tool_calls[idx]
            final_calls.append({"id": tc["id"], "type": "function",
                                "function": {"name": tc["name"], "arguments": tc["arguments"]}})
        yield ("done", {"content": "".join(content_parts), "tool_calls": final_calls})

    def _chat_baidu(self, messages):
        """百度千帆：AK/SK 换 token，非流式（备用通道，不支持 tools）"""
        token = self._baidu_token()
        model = self.cfg["chat_model"]
        url = f"https://aip.baidubce.com/rpc/2.0/ai_custom/v1/wenxinworkshop/chat/{model}?access_token={token}"
        payload = {"messages": messages, "temperature": float(self.cfg.get("temperature", 0.3))}
        try:
            r = requests.post(url, json=payload, headers={"Content-Type": "application/json"}, timeout=120)
        except requests.RequestException as e:
            raise APIError(f"百度网络错误: {e}")
        if r.status_code != 200:
            raise APIError(f"百度 API 错误 {r.status_code}: {r.text[:300]}")
        data = r.json()
        if "error_code" in data:
            raise APIError(f"百度 API 错误: {data.get('error_msg', data)}")
        content = data.get("result", "")
        if content:
            yield ("content", content)
        yield ("done", {"content": content, "tool_calls": []})

    def _chat_g4f(self, messages):
        """gpt4free 开源保底通道：逆向免费模型，无需 API Key（不支持 tools）。
        默认走实测免 Key 渠道 Yqcloud，失败后回退 g4f 全渠道轮询。"""
        if not G4F_AVAILABLE:
            raise APIError("g4f 未安装，请运行 pip install g4f")
        model = self.cfg.get("chat_model") or "gpt-4o-mini"
        # 实测当前免 Key 可用的渠道；后续 g4f 渠道变动可在此增删
        provider_names = ["Yqcloud"]
        last_err = None
        for pname in provider_names + [None]:  # None = 默认 RetryProvider 全渠道
            try:
                from g4f import Provider
                kwargs = {}
                if pname:
                    kwargs["provider"] = getattr(Provider, pname)
                client = G4FClient()
                stream = client.chat.completions.create(model=model, messages=messages,
                                                        stream=True, **kwargs)
                content_parts = []
                got_any = False
                for chunk in stream:
                    if chunk and chunk.choices and chunk.choices[0].delta and chunk.choices[0].delta.content:
                        text = chunk.choices[0].delta.content
                        content_parts.append(text)
                        got_any = True
                        yield ("content", text)
                if not got_any:
                    raise APIError(f"{pname or 'RetryProvider'} 返回空内容")
                yield ("done", {"content": "".join(content_parts), "tool_calls": []})
                return
            except Exception as e:
                last_err = e
                continue
        raise APIError(f"gpt4free 所有渠道均失败: {last_err}")

    def _baidu_token(self):
        r = requests.post("https://aip.baidubce.com/oauth/2.0/token",
                          params={"grant_type": "client_credentials",
                                  "client_id": self.cfg.get("api_key", ""),
                                  "client_secret": self.cfg.get("secret_key", "")},
                          timeout=20)
        data = r.json()
        token = data.get("access_token")
        if not token:
            raise APIError(f"百度鉴权失败: {data}")
        return token

    # ---------- 图像生成 ----------
    def generate_image(self, prompt, size="1024x1024"):
        if self.type == "baidu":
            raise APIError("百度备用通道不支持图像生成")
        url = self.cfg["base_url"].rstrip("/") + "/images/generations"
        payload = {"model": self.cfg["image_model"], "prompt": prompt, "size": size, "n": 1}
        headers = {"Authorization": f"Bearer {self.cfg['api_key']}",
                   "Content-Type": "application/json"}
        r = requests.post(url, json=payload, headers=headers, timeout=180)
        if r.status_code != 200:
            raise APIError(f"图像 API 错误 {r.status_code}: {r.text[:300]}")
        data = r.json()
        if data.get("data"):
            item = data["data"][0]
            if item.get("b64_json"):
                return base64.b64decode(item["b64_json"])
            if item.get("url"):
                img = requests.get(item["url"], timeout=60)
                img.raise_for_status()
                return img.content
        raise APIError("图像生成失败：未返回图片数据")

    # ---------- 图像理解 ----------
    def chat_vision(self, messages, image_path):
        if self.type == "baidu":
            raise APIError("百度备用通道不支持图像理解")
        with open(image_path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode()
        ext = os.path.splitext(image_path)[1].lstrip(".").lower()
        mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg", "gif": "image/gif",
                "webp": "image/webp", "bmp": "image/bmp"}.get(ext, "image/png")
        messages = messages + [{
            "role": "user",
            "content": [{"type": "text", "text": "请详细描述这张图片的内容，包括主体、场景、文字、细节等。"},
                        {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}}]
        }]
        url = self.cfg["base_url"].rstrip("/") + "/chat/completions"
        payload = {"model": self.cfg["chat_model"], "messages": messages,
                   "stream": False, "temperature": 0.5}
        headers = {"Authorization": f"Bearer {self.cfg['api_key']}",
                   "Content-Type": "application/json"}
        r = requests.post(url, json=payload, headers=headers, timeout=180)
        if r.status_code != 200:
            raise APIError(f"视觉 API 错误 {r.status_code}: {r.content.decode('utf-8', errors='replace')[:300]}")
        data = r.json()
        return data["choices"][0]["message"]["content"]


# ============ FreeLLMAPI 本地免费聚合服务（自动调度） ============
_FRELLMAPI_STATE = {"ready_at": 0.0, "key": ""}
_FRELLMAPI_BASE = "http://127.0.0.1:3001/v1"


def _http_ok(url, timeout=2.0):
    """轻量探测 HTTP 服务是否就绪"""
    try:
        r = requests.get(url, timeout=timeout)
        return r.status_code < 500
    except Exception:
        return False


def _find_freellmapi_root():
    """在常见位置寻找 FreeLLMAPI 项目根目录（含 server/server.js）"""
    cands = []
    cwd = os.getcwd()
    exe_dir = os.path.dirname(os.path.abspath(sys.argv[0]))
    code_dir = os.path.dirname(os.path.abspath(__file__))
    cands += [os.path.join(cwd, "dl", "freellmapi"),
              os.path.join(exe_dir, "dl", "freellmapi"),
              os.path.join(code_dir, "dl", "freellmapi")]
    d = code_dir
    for _ in range(6):  # 开发态向上逐级找 temp/AgnesAgent → workspace → ...
        d = os.path.dirname(d)
        cands.append(os.path.join(d, "dl", "freellmapi"))
    seen = set()
    for c in cands:
        c = os.path.abspath(c)
        if c in seen:
            continue
        seen.add(c)
        if (os.path.isfile(os.path.join(c, "server", "server.js"))
                or os.path.isfile(os.path.join(c, "server", "index.js"))
                or os.path.isfile(os.path.join(c, "server", "dist", "index.js"))):
            return c
    return None


def _freellmapi_key_from_db(root):
    """从 FreeLLMAPI sqlite settings 表读取 unified_api_key"""
    db = os.path.join(root, "server", "data", "freeapi.db")
    try:
        import sqlite3
        con = sqlite3.connect(db)
        try:
            row = con.execute("SELECT value FROM settings WHERE key='unified_api_key'").fetchone()
        finally:
            con.close()
        if row and row[0]:
            return row[0]
    except Exception:
        pass
    return ""


def _ensure_freellmapi_service(timeout=15):
    """确保本地 FreeLLMAPI Node 服务可用：探测 → 拉起 → 等待就绪。
    返回 (base_url, api_key, err)"""
    if time.time() - _FRELLMAPI_STATE["ready_at"] < 30 and _FRELLMAPI_STATE.get("key"):
        return _FRELLMAPI_BASE, _FRELLMAPI_STATE["key"], None
    if _http_ok(_FRELLMAPI_BASE.rstrip("/") + "/models", 2):
        root = _find_freellmapi_root()
        key = _freellmapi_key_from_db(root) if root else ""
        _FRELLMAPI_STATE.update({"ready_at": time.time(), "key": key})
        return _FRELLMAPI_BASE, key, None
    root = _find_freellmapi_root()
    if not root:
        return _FRELLMAPI_BASE, "", "未找到 FreeLLMAPI 目录（dl/freellmapi），跳过免费聚合"
    pj = os.path.join(root, "server")
    if os.path.isfile(os.path.join(pj, "server.js")):
        script = "server.js"
    elif os.path.isfile(os.path.join(pj, "index.js")):
        script = "index.js"
    elif os.path.isfile(os.path.join(pj, "dist", "index.js")):
        script = os.path.join("dist", "index.js")
    else:
        return _FRELLMAPI_BASE, "", "未找到 FreeLLMAPI 启动脚本（server.js / index.js / dist/index.js）"
    log_path = os.path.join(pj, "freellmapi.log")
    log = open(log_path, "a", encoding="utf-8")
    flags = 0x08000000 if os.name == "nt" else 0  # CREATE_NO_WINDOW
    try:
        subprocess.Popen(["node", script], cwd=pj, stdout=log, stderr=log, creationflags=flags)
    except Exception as e:
        return _FRELLMAPI_BASE, "", f"启动 FreeLLMAPI 失败: {e}"
    deadline = time.time() + timeout
    while time.time() < deadline:
        time.sleep(0.8)
        if _http_ok(_FRELLMAPI_BASE.rstrip("/") + "/models", 1):
            key = _freellmapi_key_from_db(root)
            _FRELLMAPI_STATE.update({"ready_at": time.time(), "key": key})
            return _FRELLMAPI_BASE, key, None
    return _FRELLMAPI_BASE, "", "FreeLLMAPI 服务启动超时"


# ============ ComfyUI 本地视频（Wan2.1 GGUF）自动调度 ============
COMFYUI_HOST = "127.0.0.1"
COMFYUI_PORT = 8188
COMFYUI_BASE = f"http://{COMFYUI_HOST}:{COMFYUI_PORT}"


def _locate_comfyui():
    """定位 ComfyUI 目录与可用的 python 解释器。
    返回 (comfyui_dir, python_exe, err)"""
    exe_dir = os.path.dirname(os.path.abspath(sys.argv[0]))
    code_dir = os.path.dirname(os.path.abspath(__file__))
    cand_dirs = [
        os.path.join(exe_dir, "comfyui"),
        os.path.join(code_dir, "comfyui"),
    ]
    d = code_dir
    for _ in range(6):  # 开发态：向上找 v15_repro/AgnesAgent_v15_onedir/comfyui 等
        d = os.path.dirname(d)
        cand_dirs.append(os.path.join(d, "comfyui"))
        cand_dirs.append(os.path.join(d, "AgnesAgent_v15_onedir", "comfyui"))
        cand_dirs.append(os.path.join(d, "v15_repro", "AgnesAgent_v15_onedir", "comfyui"))
    comfyui_dir = None
    for c in cand_dirs:
        c = os.path.abspath(c)
        if os.path.isfile(os.path.join(c, "main.py")):
            comfyui_dir = c
            break
    if not comfyui_dir:
        return None, None, "未找到 ComfyUI（需含 main.py 的 comfyui 目录）"
    # python 解释器候选
    cand_py = [
        os.path.join(os.path.dirname(comfyui_dir), "comfy_venv", "Scripts", "python.exe"),
        os.path.join(os.path.dirname(comfyui_dir), "venv", "Scripts", "python.exe"),
        os.path.join(comfyui_dir, "python_embeded", "python.exe"),
        os.path.join(comfyui_dir, "venv", "Scripts", "python.exe"),
        sys.executable,
    ]
    for p in cand_py:
        if p and os.path.isfile(p):
            return comfyui_dir, p, None
    return comfyui_dir, sys.executable, None


def _comfyui_running():
    return _http_ok(COMFYUI_BASE + "/system_stats", 2)


def _ensure_comfyui_running(timeout=180):
    """确保 ComfyUI 已启动（纯 CPU：main.py --cpu --port 8188），
    返回 (ok, msg)"""
    if _comfyui_running():
        return True, "ComfyUI 已在运行"
    comfyui_dir, py, err = _locate_comfyui()
    if err:
        return False, err
    log_path = os.path.join(comfyui_dir, "agnes_comfyui.log")
    log = open(log_path, "a", encoding="utf-8")
    flags = 0x08000000 if os.name == "nt" else 0
    try:
        subprocess.Popen([py, "main.py", "--cpu", "--port", str(COMFYUI_PORT)],
                         cwd=comfyui_dir, stdout=log, stderr=log, creationflags=flags)
    except Exception as e:
        return False, f"启动 ComfyUI 失败: {e}"
    deadline = time.time() + timeout
    while time.time() < deadline:
        time.sleep(5)
        if _comfyui_running():
            return True, "ComfyUI 已自动启动（纯 CPU 模式）"
    return False, f"ComfyUI 启动超时（{timeout}s），日志见 {log_path}"


def _build_wan_workflow(prompt, unet_name="wan13b_q5.gguf",
                        clip_name="umt5-xxl-Q3_K_M.gguf",
                        vae_name="Wan2_1_VAE_bf16.safetensors",
                        length=81, width=832, height=480, steps=20, cfg=4.0):
    """构造 Wan2.1 T2V GGUF ComfyUI 工作流（WanVideoWrapper 兼容节点）"""
    return {
        "1": {"class_type": "UnetLoaderGGUF", "inputs": {"unet_name": unet_name}},
        "2": {"class_type": "CLIPLoaderGGUF", "inputs": {"clip_name": clip_name, "type": "wan"}},
        "3": {"class_type": "VAELoader", "inputs": {"vae_name": vae_name}},
        "4": {"class_type": "CLIPTextEncode", "inputs": {"text": prompt, "clip": ["2", 0]}},
        "5": {"class_type": "CLIPTextEncode", "inputs": {"text": "", "clip": ["2", 0]}},
        "6": {"class_type": "WanImageToVideo", "inputs": {
            "width": width, "height": height, "length": length, "batch_size": 1,
            "positive": ["4", 0], "negative": ["5", 0],
            "vae": ["3", 0]}},
        "7": {"class_type": "KSampler", "inputs": {
            "seed": random.randint(0, 2 ** 31 - 1), "steps": steps, "cfg": cfg,
            "sampler_name": "euler", "scheduler": "simple", "denoise": 1.0,
            "model": ["1", 0], "positive": ["6", 0], "negative": ["6", 1], "latent_image": ["6", 2]}},
        "8": {"class_type": "VAEDecode", "inputs": {"samples": ["7", 0], "vae": ["3", 0]}},
        "9": {"class_type": "SaveWEBM", "inputs": {
            "images": ["8", 0], "filename_prefix": "agnes_wan",
            "codec": "vp9", "fps": 16, "crf": 42}},
    }


def _generate_video_local_blocking(prompt, progress_cb=None, length=81):
    """调用本地 ComfyUI 生成 Wan2.1 视频（自动定位/启动 ComfyUI）。
    返回 (video_path, msg)；失败返回 (None, err)"""
    if progress_cb:
        progress_cb("正在定位 ComfyUI…")
    comfyui_dir, py, err = _locate_comfyui()
    if err:
        return None, err
    ok, msg = _ensure_comfyui_running()
    if not ok:
        return None, msg
    if progress_cb:
        progress_cb("ComfyUI 就绪，提交 Wan2.1 生成任务…")
    workflow = _build_wan_workflow(prompt, length=length)
    try:
        r = requests.post(COMFYUI_BASE + "/prompt", json={"prompt": workflow}, timeout=30)
        if r.status_code != 200:
            return None, f"ComfyUI 提交失败 {r.status_code}: {r.text[:200]}"
        prompt_id = r.json()["prompt_id"]
    except Exception as e:
        return None, f"ComfyUI 提交异常: {e}"
    # 轮询 /history 直到完成（CPU 生成 5 秒视频约 20-40 分钟）
    deadline = time.time() + 60 * 50
    last_status = ""
    while time.time() < deadline:
        time.sleep(10)
        try:
            h = requests.get(f"{COMFYUI_BASE}/history/{prompt_id}", timeout=10).json()
        except Exception:
            continue
        hist = h.get(prompt_id)
        if not hist:
            continue
        if hist.get("status", {}).get("status_str") == "error":
            errs = hist["status"].get("messages", [])
            return None, f"ComfyUI 执行出错: {json.dumps(errs, ensure_ascii=False)[:300]}"
        if hist.get("outputs"):
            for node_id, out in hist["outputs"].items():
                for v in out.get("gifs", []) + out.get("videos", []) + out.get("images", []):
                    fname = v.get("filename")
                    if not fname:
                        continue
                    try:
                        vr = requests.get(COMFYUI_BASE + "/view", params={
                            "filename": fname, "subfolder": v.get("subfolder", ""),
                            "type": v.get("type", "output")}, timeout=30)
                        if vr.status_code == 200:
                            video_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "videos")
                            os.makedirs(video_dir, exist_ok=True)
                            ext = os.path.splitext(fname)[1].lower() or ".mp4"
                            out_path = os.path.join(video_dir, f"video_local_{time.strftime('%Y%m%d_%H%M%S')}{ext}")
                            with open(out_path, "wb") as f:
                                f.write(vr.content)
                            return out_path, "本地 Wan2.1 视频生成完成"
                    except Exception:
                        continue
        # 进度提示（每 2 分钟报一次）
        status = hist.get("status", {})
        running = status.get("running")
        if running and running != last_status:
            last_status = running
            if progress_cb:
                progress_cb(f"正在推理…（CPU 生成约需 20-40 分钟，当前已运行 {int((time.time() - (deadline - 60*50)) / 60)} 分钟）")
    return None, "本地视频生成超时（50 分钟），请检查 ComfyUI 日志"


def build_clients(config, mode="cloud"):
    """按配置顺序返回启用的 Provider 客户端列表。

    物理隔离（重要）：本地与在线严格二选一，绝不互相兜底切换——
      mode=local：只返回本地 ollama/llama-server，绝不包含任何在线 API；
      mode=cloud：只返回在线 provider，绝不包含本地 ollama，也不拉起本地服务。
    从根上杜绝在线 API 话术污染本地 AI、以及本地失败静默降级到云端的问题。"""
    clients = []
    providers = dict(config.get("providers", {}))
    temp = float(config.get("temperature", 0.3))
    if mode == "local":
        # ---- 本地模式：物理隔离，只保留本地模型 ----
        oll = providers.get("ollama")
        if oll and oll.get("enabled") and _start_llama_server():
            oll = dict(oll)
            oll["base_url"] = _LLAMA_BASE_URL
            oll["chat_model"] = local_llm_model_name()
            oll.setdefault("temperature", temp)
            clients.append(ProviderClient("ollama", oll))
        return clients
    # ---- 云端模式：物理隔离，绝不包含本地 ollama，也不拉起本地服务 ----
    # FreeLLMAPI 自动调度：启用时自动拉起本地服务并注入 key
    fre = providers.get("freellmapi")
    if fre and fre.get("enabled"):
        base, key, ferr = _ensure_freellmapi_service()
        if key:
            fre = dict(fre)
            fre["api_key"] = key
            if not fre.get("base_url"):
                fre["base_url"] = base
            providers["freellmapi"] = fre
        elif ferr:
            pass  # 拉不起来则跳过，其余 provider 照常兜底
    order = ("agnes", "freellmapi", "zhipu", "baidu", "g4f")
    # 主 provider 优先，其余按序；g4f 免 Key 恒兜底
    for pid in order:
        p = providers.get(pid)
        if not (p and p.get("enabled")):
            continue
        if pid == "g4f" or p.get("api_key", "").strip():
            p = dict(p)
            p.setdefault("temperature", temp)
            clients.append(ProviderClient(pid, p))
    # 用户自定义 provider
    for pid, p in providers.items():
        if pid in ("agnes", "zhipu", "baidu", "ollama", "g4f", "freellmapi"):
            continue
        if p and p.get("enabled") and p.get("api_key", "").strip():
            p = dict(p)
            p.setdefault("temperature", temp)
            clients.append(ProviderClient(pid, p))
    if not clients:
        # 至少保留主 provider（即使 key 为空也尝试，报错可见）
        agnes = providers.get("agnes") or DEFAULT_CONFIG["providers"]["agnes"]
        agnes = dict(agnes)
        agnes.setdefault("temperature", temp)
        clients.append(ProviderClient("agnes", agnes))
    return clients


def pick_image_client(clients):
    """选择可生成图像的 provider（有 image_model 且类型非 baidu）"""
    for c in clients:
        if c.cfg.get("image_model"):
            return c
    return None


# ============ 工作线程 ============
class ChatWorker(QThread):
    """对话 + 工具调用循环，支持多 Provider 自动兜底"""
    token = Signal(str)            # 流式 token
    tool_used = Signal(str)        # 工具/切换提示
    tool_result = Signal(dict)     # 结构化工具结果（用于 UI 卡片渲染）
    finished = Signal(str)         # 最终完整内容
    error = Signal(str)

    def __init__(self, clients, messages, tools=None, parent=None):
        super().__init__(parent)
        self.clients = clients
        self.messages = messages
        self.tools = tools if tools is not None else TOOLS
        self._any_token = False
        self._full_content = ""

    def run(self):
        err = None
        for i, client in enumerate(self.clients):
            self._any_token = False
            self._full_content = ""
            try:
                self._dialog_loop(client)
                self.finished.emit(self._full_content)
                return
            except APIError as e:
                err = str(e)
                # 已输出内容则不切换（避免重复），直接报错
                if self._any_token:
                    self.error.emit(err)
                    return
                if i < len(self.clients) - 1:
                    self.tool_used.emit(f"⚠ {client.name} 不可用（{err}），已自动切换到 {self.clients[i+1].name}")
            except Exception as e:
                err = str(e)
                if self._any_token:
                    self.error.emit(err)
                    return
                if i < len(self.clients) - 1:
                    self.tool_used.emit(f"⚠ {client.name} 出错（{err}），已自动切换到 {self.clients[i+1].name}")
                self.error.emit(str(err or "所有 API 均不可用"))

    def _sanitize_messages(self, client):
        """非视觉 Provider 或非视觉模型，把多模态消息中的图片块降级为纯文本占位"""
        if client.type != "openai":
            return self._strip_image_blocks(self.messages)
        model = str(client.cfg.get("chat_model", "")).lower()
        vision_models = ("agnes", "vision", "glm-4v", "qwen-vl", "gpt-4o", "gpt-4.1", "claude", "gemini", "gemma")
        if any(v in model for v in vision_models):
            return self.messages
        msgs = self._strip_image_blocks(self.messages)
        # 本地模型（包内 Qwen3.6-35B）CPU 推理 prefill 慢，截断历史控制单次等待时长
        if "18080" in str(client.cfg.get("base_url", "")) or "qwen3.6" in model:
            return self._truncate_for_local(msgs)
        return msgs

    @staticmethod
    def _truncate_for_local(messages, max_chars=4500):
        """保留 system 与最近对话，按字符粗估 token 截断（中文 1 字≈1 token），控制 prefill 时长"""
        sys_msgs = [m for m in messages if m.get("role") == "system"]
        rest = [m for m in messages if m.get("role") != "system"]
        budget = max_chars - sum(len(str(m.get("content", ""))) for m in sys_msgs)
        if budget <= 0:
            return messages[-2:]
        keep = []
        used = 0
        for m in reversed(rest):
            length = len(str(m.get("content", "")))
            if used + length > budget and keep:
                break
            keep.append(m)
            used += length
        keep.reverse()
        return sys_msgs + keep

    @staticmethod
    def _strip_image_blocks(messages):
        msgs = []
        for m in messages:
            if isinstance(m.get("content"), list):
                parts = []
                for p in m["content"]:
                    if isinstance(p, dict) and p.get("type") == "text":
                        parts.append(p.get("text", ""))
                    elif isinstance(p, dict) and p.get("type") == "image_url":
                        parts.append("[图片]")
                m2 = dict(m)
                m2["content"] = "\n".join(parts) or None
                msgs.append(m2)
            else:
                msgs.append(m)
        return msgs

    def _dialog_loop(self, client):
        loop_count = 0
        while loop_count < 6:
            loop_count += 1
            got_tool = False
            for kind, data in client.chat_stream(self._sanitize_messages(client), tools=self.tools):
                if kind == "content":
                    self._any_token = True
                    self.token.emit(data)
                    self._full_content += data
                elif kind == "done":
                    tc = data.get("tool_calls", [])
                    if tc:
                        got_tool = True
                        assistant_msg = {"role": "assistant",
                                         "content": data.get("content") or None}
                        calls = []
                        for c in tc:
                            calls.append({"id": c["id"], "type": "function",
                                          "function": {"name": c["function"]["name"],
                                                       "arguments": c["function"]["arguments"]}})
                        assistant_msg["tool_calls"] = calls
                        self.messages.append(assistant_msg)
                        for c in calls:
                            name = c["function"]["name"]
                            try:
                                args = json.loads(c["function"]["arguments"] or "{}")
                            except Exception:
                                args = {}
                            self.tool_used.emit(f"[工具] {name}")
                            result = dispatch_tool(name, args)
                            self.messages.append({"role": "tool", "tool_call_id": c["id"],
                                                  "content": json.dumps(result, ensure_ascii=False)})
                            self.tool_result.emit(result)
                    break
            if not got_tool:
                break


def _locate_sd_engine():
    """定位 v15 自带 sd.cpp 本地生图引擎；优先查打包资源(_MEIPASS)，其次查 exe 同目录外部资源；未找到返回 (None, 原因)"""
    cands = []
    if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
        cands.append(sys._MEIPASS)
    cands.append(os.path.dirname(sys.executable) if getattr(sys, "frozen", False)
                 else os.path.dirname(os.path.abspath(__file__)))
    for base in cands:
        exe = os.path.join(base, "dl", "sd_cpp", "sd-cli.exe")
        models_dir = os.path.join(base, "dl", "sd_models")
        if not os.path.exists(exe):
            continue
        models = []
        if os.path.isdir(models_dir):
            models = [f for f in os.listdir(models_dir)
                      if f.lower().endswith((".safetensors", ".gguf", ".ckpt"))]
        if not models:
            continue
        return {"exe": exe, "model": os.path.join(models_dir, models[0])}, None
    return None, "未找到本地生图引擎 dl/sd_cpp/sd-cli.exe（需 v15 完整目录，且 exe 同级需含 dl/sd_cpp 与 dl/sd_models）"


def generate_image_local_sd(prompt, save_path, width=256, height=256, steps=20, cfg_scale=7):
    """调用 v15 内置 sd.cpp 引擎本地生成图片（纯 CPU 离线，约 1-2 分钟/张）"""
    info, err = _locate_sd_engine()
    if err:
        raise RuntimeError(err)
    prompt_file = save_path + ".prompt.txt"
    with open(prompt_file, "w", encoding="utf-8") as f:
        f.write(prompt)
    try:
        proc = subprocess.run(
            [info["exe"], "-m", info["model"], "--prompt-file", prompt_file,
             "-o", save_path, "--cfg-scale", str(cfg_scale), "--steps", str(steps),
             "--width", str(width), "--height", str(height)],
            capture_output=True, text=True, timeout=900, encoding="utf-8", errors="replace",
        )
    finally:
        if os.path.exists(prompt_file):
            try:
                os.remove(prompt_file)
            except Exception:
                pass
    if proc.returncode != 0:
        raise RuntimeError(f"sd-cli 退出码 {proc.returncode}：{(proc.stderr or proc.stdout)[-400:]}")
    if not os.path.exists(save_path) or os.path.getsize(save_path) < 500:
        raise RuntimeError("sd-cli 未产出图片文件")
    return save_path


class LocalImageWorker(QThread):
    """本地生图：sd.cpp + Realistic Vision V6（纯 CPU 离线）"""
    done = Signal(str)
    error = Signal(str)

    def __init__(self, prompt, save_dir, parent=None):
        super().__init__(parent)
        self.prompt = prompt
        self.save_dir = save_dir

    def run(self):
        os.makedirs(self.save_dir, exist_ok=True)
        path = os.path.join(self.save_dir, f"img_local_{time.strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:6]}.png")
        try:
            generate_image_local_sd(self.prompt, path)
            self.done.emit(path)
        except Exception as e:
            self.error.emit(str(e))


class ImageWorker(QThread):
    done = Signal(str)
    error = Signal(str)

    def __init__(self, client, prompt, save_dir, parent=None):
        super().__init__(parent)
        self.client = client
        self.prompt = prompt
        self.save_dir = save_dir

    def run(self):
        os.makedirs(self.save_dir, exist_ok=True)
        path = os.path.join(self.save_dir, f"image_{time.strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:6]}.png")
        try:
            img_bytes = self.client.generate_image(self.prompt)
            with open(path, "wb") as f:
                f.write(img_bytes)
            self.done.emit(path)
        except Exception as e:
            # 主 Provider 失败 → Pollinations 免 Key 兜底（已实测可用）
            try:
                pollinations_image(self.prompt, path)
                if os.path.exists(path) and os.path.getsize(path) > 500:
                    self.done.emit(path)
                    return
                self.error.emit(str(e))
            except Exception as e2:
                self.error.emit(f"主通道失败：{e}；兜底通道失败：{e2}")


class VoiceRecWorker(QThread):
    """语音输入：麦克风录音（最长12秒，可提前停止）→ 免费 ASR 转文字"""
    done = Signal(str)
    error = Signal(str)

    def __init__(self, parent=None):
        super().__init__(parent)
        self._stop = False

    def stop(self):
        self._stop = True

    def run(self):
        if not _SOUNDDEVICE_OK:
            self.error.emit("未安装 sounddevice，请运行 pip install sounddevice")
            return
        try:
            import numpy as np
            import wave as _wave
        except Exception as e:
            self.error.emit(f"缺少录音依赖 numpy/wave: {e}")
            return
        sample_rate = 16000
        frames = []
        try:
            def cb(indata, frames_ok, time_info, status):
                frames.append(indata.copy())
            stream = sd.InputStream(samplerate=sample_rate, channels=1, dtype="int16", callback=cb)
            stream.start()
            deadline = time.time() + 12
            while not self._stop and time.time() < deadline:
                time.sleep(0.1)
            stream.stop()
            stream.close()
        except Exception as e:
            self.error.emit(f"录音失败: {e}")
            return
        if not frames:
            self.error.emit("没有录到声音，请靠近麦克风重试")
            return
        try:
            data = np.concatenate(frames).reshape(-1)
            os.makedirs(TTS_CACHE_DIR, exist_ok=True)
            path = os.path.join(TTS_CACHE_DIR, f"voice_{time.strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:6]}.wav")
            with _wave.open(path, "wb") as wf:
                wf.setnchannels(1)
                wf.setsampwidth(2)
                wf.setframerate(sample_rate)
                wf.writeframes(data.tobytes())
        except Exception as e:
            self.error.emit(f"保存录音失败: {e}")
            return
        text, err = transcribe_audio(path)
        if err:
            self.error.emit(err)
            return
        if not text.strip():
            self.error.emit("未能识别出语音内容")
            return
        self.done.emit(text.strip())


class VisionWorker(QThread):
    done = Signal(str)
    error = Signal(str)

    def __init__(self, client, image_path, parent=None):
        super().__init__(parent)
        self.client = client
        self.image_path = image_path

    def run(self):
        try:
            result = self.client.chat_vision([{"role": "user", "content": "请分析这张图片"}], self.image_path)
            self.done.emit(result)
        except Exception as e:
            self.error.emit(str(e))


class TTSWorker(QThread):
    done = Signal(str)   # mp3 路径

    def __init__(self, text, parent=None):
        super().__init__(parent)
        self.text = text

    def run(self):
        try:
            import edge_tts
            import asyncio
            path = os.path.join(TTS_CACHE_DIR, f"tts_{time.strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:6]}.mp3")

            async def _tts():
                communicate = edge_tts.Communicate(self.text, "zh-CN-XiaoxiaoNeural", rate="+0%")
                await communicate.save(path)

            asyncio.run(_tts())
            if os.path.getsize(path) < 500:
                raise Exception("edge-tts 返回空文件")
            self.done.emit(path)
            # 清理旧缓存，最多保留 30 个
            try:
                files = sorted(os.listdir(TTS_CACHE_DIR))
                for f in files[:-30]:
                    os.remove(os.path.join(TTS_CACHE_DIR, f))
            except Exception:
                pass
            return
        except Exception as e:
            # edge-tts 失败，降级本地 TTS（sherpa-onnx vits-zh 离线）
            try:
                local_path = os.path.join(TTS_CACHE_DIR, f"tts_local_{time.strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:6]}.wav")
                ok, msg = local_tts(self.text, local_path)
                if ok:
                    self.done.emit(local_path)
                    return
            except Exception:
                pass
            self.done.emit("")


class VideoWorker(QThread):
    """文生视频：AgnesAI API → 自动切本地 ComfyUI Wan2.1。期间持续发进度信号"""
    done = Signal(str, str)   # (path, msg)
    error = Signal(str)
    progress = Signal(str)    # 状态文本（等待计时在 UI 侧每 10s 刷新一次，这里仅发阶段切换）

    def __init__(self, prompt, parent=None):
        super().__init__(parent)
        self.prompt = prompt

    def run(self):
        try:
            self.progress.emit("正在连接 AgnesAI 视频服务…")
            path, err = _generate_video_api(self.prompt, 5)
            if path:
                self.done.emit(path, "视频已生成（AgnesAI 免费 API）")
                return
            self.progress.emit(f"云端失败（{err}），自动切换本地 ComfyUI + Wan2.1…")
            path, msg = _generate_video_local(self.prompt, 5)
            if path:
                self.done.emit(path, msg)
                return
            self.error.emit(f"云端与本地均失败：{err} | {msg}")
        except Exception as e:
            self.error.emit(f"视频生成异常：{e}")


class LocalVideoWorker(QThread):
    """本地生视频：ComfyUI + Wan2.1 GGUF 自动调度，带真实进度回调"""
    done = Signal(str, str)   # (path, msg)
    error = Signal(str)
    progress = Signal(str)

    def __init__(self, prompt, seconds=5, parent=None):
        super().__init__(parent)
        self.prompt = prompt
        self.seconds = seconds

    def run(self):
        try:
            length = 81 if self.seconds <= 6 else 121
            path, msg = _generate_video_local_blocking(
                self.prompt, progress_cb=lambda s: self.progress.emit(s), length=length)
            if path:
                self.done.emit(path, msg)
            else:
                self.error.emit(msg)
        except Exception as e:
            self.error.emit(f"本地生视频异常：{e}")


# ============ 模型市场（Wan2.1 系列模型一键下载） ============
MODEL_MARKET = [
    {"name": "Wan2.1-1.3B T2V（官方 fp16）", "file": "wan2.1_t2v_1.3B_bf16.safetensors", "subdir": "diffusion_models",
     "url": "https://huggingface.co/Comfy-Org/Wan_2.1_ComfyUI_repackaged/resolve/main/split_files/diffusion_models/wan2.1_t2v_1.3B_bf16.safetensors",
     "size": "约 2.6GB", "desc": "轻量 1.3B 文生视频，CPU 可跑，画质优于 GGUF Q5"},
    {"name": "Wan2.1-14B T2V Q4_K_S（GGUF）", "file": "wan2.1-t2v-14b-Q4_K_S.gguf", "subdir": "unet",
     "url": "https://huggingface.co/city96/Wan2.1-T2V-14B-gguf/resolve/main/wan2.1-t2v-14b-Q4_K_S.gguf",
     "size": "约 10GB", "desc": "14B 高画质（Q4 量化），内存 12GB+ 可试"},
    {"name": "Wan2.1-14B T2V Q5_K_S（GGUF）", "file": "wan2.1-t2v-14b-Q5_K_S.gguf", "subdir": "unet",
     "url": "https://huggingface.co/city96/Wan2.1-T2V-14B-gguf/resolve/main/wan2.1-t2v-14b-Q5_K_S.gguf",
     "size": "约 12GB", "desc": "14B Q5 平衡画质与体积，内存 16GB+ 推荐"},
    {"name": "umt5-xxl 文本编码器（fp8）", "file": "umt5-xxl-enc-fp8_e4m3fn.safetensors", "subdir": "text_encoders",
     "url": "https://huggingface.co/Kijai/WanVideo_comfy/resolve/main/umt5-xxl-enc-fp8_e4m3fn.safetensors",
     "size": "约 6GB", "desc": "高精度文本编码，可替换 umt5-xxl-Q3_K_M.gguf 提升语义理解"},
    {"name": "Wan2.1 VAE（bf16）", "file": "Wan2_1_VAE_bf16.safetensors", "subdir": "vae",
     "url": "https://huggingface.co/Kijai/WanVideo_comfy/resolve/main/Wan2_1_VAE_bf16.safetensors",
     "size": "约 0.25GB", "desc": "官方 VAE，画质更稳，可替换 wan21_vae.safetensors"},
]


def _comfyui_models_root():
    """返回 ComfyUI models 根目录（不存在则 None）"""
    comfyui_dir, _, err = _locate_comfyui()
    if err or not comfyui_dir:
        return None, err
    root = os.path.join(comfyui_dir, "models")
    os.makedirs(root, exist_ok=True)
    return root, None


class ModelDownloadWorker(QThread):
    """下载单个模型到 ComfyUI models 子目录，带进度与取消"""
    progress = Signal(int, int, str)   # (当前MB, 总MB, 文本)
    done = Signal(str, str)            # (msg, err) 二选一
    finished_one = Signal(str)         # 完成条目名

    def __init__(self, entries, models_root, parent=None):
        super().__init__(parent)
        self.entries = entries
        self.models_root = models_root
        self._cancel = False

    def cancel(self):
        self._cancel = True

    def run(self):
        ok_list, err_list = [], []
        for e in self.entries:
            if self._cancel:
                err_list.append(f"{e['name']}（已取消）")
                break
            sub = os.path.join(self.models_root, e["subdir"])
            os.makedirs(sub, exist_ok=True)
            dst = os.path.join(sub, e["file"])
            if os.path.exists(dst) and os.path.getsize(dst) > 1024 * 1024:
                self.progress.emit(0, 0, f"已存在，跳过：{e['file']}")
                ok_list.append(e["name"])
                self.finished_one.emit(e["name"])
                continue
            tmp = dst + ".part"
            try:
                self.progress.emit(0, 0, f"开始下载：{e['file']}")
                with requests.get(e["url"], stream=True, timeout=60) as r:
                    r.raise_for_status()
                    total = int(r.headers.get("Content-Length", 0) or 0)
                    done_b = 0
                    with open(tmp, "wb") as f:
                        for chunk in r.iter_content(chunk_size=1024 * 256):
                            if self._cancel:
                                raise RuntimeError("cancelled")
                            if chunk:
                                f.write(chunk)
                                done_b += len(chunk)
                                if total:
                                    self.progress.emit(int(done_b // (1024 * 1024)),
                                                       int(total // (1024 * 1024)),
                                                       f"下载中：{e['file']} {done_b / 1048576:.0f}/{total / 1048576:.0f} MB")
                os.replace(tmp, dst)
                ok_list.append(e["name"])
                self.finished_one.emit(e["name"])
                self.progress.emit(0, 0, f"完成：{e['file']}")
            except Exception as ex:
                err_list.append(f"{e['name']}: {ex}")
                try:
                    os.remove(tmp)
                except Exception:
                    pass
        if err_list:
            self.done.emit("", "\n".join(err_list))
        else:
            self.done.emit(f"下载完成：{', '.join(ok_list)}", "")


class ModelMarketDialog(QDialog):
    """模型市场：列表 + 勾选下载 + 进度"""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("模型市场 - Wan2.1 系列")
        self.resize(760, 520)
        root, err = _comfyui_models_root()
        self.models_root = root
        layout = QVBoxLayout(self)
        tip = QLabel(f"下载目标：{root or err or '未找到 ComfyUI'}")
        tip.setWordWrap(True)
        layout.addWidget(tip)
        self.table = QTableWidget(len(MODEL_MARKET), 5)
        self.table.setHorizontalHeaderLabels(["选择", "模型", "体积", "说明", "状态"])
        self.table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeToContents)
        self.table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeToContents)
        self.table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeToContents)
        self.table.horizontalHeader().setSectionResizeMode(3, QHeaderView.Stretch)
        self.table.horizontalHeader().setSectionResizeMode(4, QHeaderView.ResizeToContents)
        self.table.verticalHeader().setVisible(False)
        self._checkboxes = []
        self._status_labels = []
        for i, e in enumerate(MODEL_MARKET):
            cb = QCheckBox()
            self._checkboxes.append(cb)
            cell = QWidget()
            lay = QHBoxLayout(cell)
            lay.setContentsMargins(4, 0, 4, 0)
            lay.addWidget(cb)
            self.table.setCellWidget(i, 0, cell)
            self.table.setItem(i, 1, QTableWidgetItem(e["name"]))
            self.table.setItem(i, 2, QTableWidgetItem(e["size"]))
            self.table.setItem(i, 3, QTableWidgetItem(e["desc"]))
            st = QLabel("未下载")
            self._status_labels.append(st)
            self.table.setCellWidget(i, 4, st)
        layout.addWidget(self.table)
        self.progress_bar = QLabel("")
        self.progress_bar.setWordWrap(True)
        layout.addWidget(self.progress_bar)
        btns = QHBoxLayout()
        self.btn_download = QPushButton("下载选中")
        self.btn_cancel = QPushButton("取消下载")
        self.btn_close = QPushButton("关闭")
        self.btn_download.clicked.connect(self.start_download)
        self.btn_cancel.clicked.connect(self.cancel_download)
        self.btn_close.clicked.connect(self.reject)
        btns.addWidget(self.btn_download)
        btns.addWidget(self.btn_cancel)
        btns.addStretch(1)
        btns.addWidget(self.btn_close)
        layout.addLayout(btns)
        self._worker = None
        self.refresh_status()

    def refresh_status(self):
        for i, e in enumerate(MODEL_MARKET):
            if self.models_root:
                p = os.path.join(self.models_root, e["subdir"], e["file"])
                self._status_labels[i].setText("已存在" if os.path.exists(p) else "未下载")
            else:
                self._status_labels[i].setText("—")

    def start_download(self):
        if self._worker and self._worker.isRunning():
            return
        if not self.models_root:
            QMessageBox.warning(self, "提示", "未找到 ComfyUI，无法下载模型")
            return
        entries = [e for i, e in enumerate(MODEL_MARKET) if self._checkboxes[i].isChecked()]
        if not entries:
            QMessageBox.information(self, "提示", "请先勾选要下载的模型")
            return
        self.btn_download.setEnabled(False)
        self.progress_bar.setText("准备下载…")
        self._worker = ModelDownloadWorker(entries, self.models_root, self)
        self._worker.progress.connect(self.on_dl_progress)
        self._worker.finished_one.connect(self.on_dl_one)
        self._worker.done.connect(self.on_dl_done)
        self._worker.start()

    def cancel_download(self):
        if self._worker and self._worker.isRunning():
            self._worker.cancel()
            self.progress_bar.setText("正在取消…")

    def on_dl_progress(self, cur, total, text):
        self.progress_bar.setText(text)

    def on_dl_one(self, name):
        self.refresh_status()

    def on_dl_done(self, msg, err):
        self.btn_download.setEnabled(True)
        if err:
            self.progress_bar.setText(f"部分失败：\n{err}")
            QMessageBox.warning(self, "下载完成（有失败）", err)
        else:
            self.progress_bar.setText(msg)
            self.refresh_status()
            QMessageBox.information(self, "下载完成", msg)


class ImageDownloadWorker(QThread):
    """下载图片缩略图到本地缓存（供图片网格渲染）"""
    done = Signal(list)  # [(local_path, thumb_url, orig_url, title)]

    def __init__(self, items, parent=None):
        super().__init__(parent)
        self.items = items

    def run(self):
        saved = []
        for it in self.items:
            thumb = it.get("thumb") or it.get("url", "")
            orig = it.get("url", "")
            title = it.get("title", "")
            if not thumb or not thumb.startswith("http"):
                continue
            h = hashlib.md5(thumb.encode()).hexdigest()[:16]
            local = os.path.join(SEARCH_CACHE_DIR, f"{h}.jpg")
            if not os.path.exists(local):
                try:
                    r = requests.get(thumb, headers={**HEADERS, "Referer": "https://www.bing.com/"}, timeout=15)
                    if r.status_code == 200 and len(r.content) > 500:
                        with open(local, "wb") as f:
                            f.write(r.content)
                except Exception:
                    continue
            if os.path.exists(local):
                saved.append((local, thumb, orig, title))
            if len(saved) >= 12:
                break
        self.done.emit(saved)


# ============ 微信/QQ 风格消息组件 ============
class AvatarLabel(QLabel):
    """圆形头像"""

    def __init__(self, text, color, size=40, parent=None):
        super().__init__(text, parent)
        self.setFixedSize(size, size)
        self.setAlignment(Qt.AlignCenter)
        self.setStyleSheet(
            f"QLabel {{ background: {color}; color: #fff; border-radius: {size // 2}px;"
            f" font-size: {size // 2 - 4}px; font-weight: bold; }}")
        if FLUENT_OK:
            eff = QGraphicsDropShadowEffect(self)
            eff.setBlurRadius(8)
            eff.setOffset(0, 2)
            eff.setColor(QColor(0, 0, 0, 40))
            self.setGraphicsEffect(eff)


class MessageRow(QWidget):
    """一行消息：头像 + (时间 + 气泡)，微信式左右分布"""

    def __init__(self, who="ai", parent=None):
        super().__init__(parent)
        self.who = who
        root = QHBoxLayout(self)
        root.setContentsMargins(4, 2, 4, 2)
        root.setSpacing(10)
        self.avatar = AvatarLabel("AI", "#7c3aed") if who == "ai" else AvatarLabel("我", "#2b5bff")
        self.time_label = QLabel("")
        self.time_label.setObjectName("msgTime")
        col = QVBoxLayout()
        col.setSpacing(5)
        if who == "ai":
            col.addWidget(self.time_label, 0, Qt.AlignLeft)
            self.bubble = BubbleFrame("ai")
            col.addWidget(self.bubble, 0, Qt.AlignLeft)
            root.addWidget(self.avatar, 0, Qt.AlignTop)
            root.addLayout(col)
            root.addStretch()
        else:
            col.addWidget(self.time_label, 0, Qt.AlignRight)
            self.bubble = BubbleFrame("user")
            col.addWidget(self.bubble, 0, Qt.AlignRight)
            root.addStretch()
            root.addLayout(col)
            root.addWidget(self.avatar, 0, Qt.AlignTop)

    def set_time(self, ts):
        self.time_label.setText(ts)


def current_time_str():
    return time.strftime("%H:%M")


def _fmt_session_time(ts):
    """会话列表时间：今天显示 HH:MM，昨天显示「昨天」，更早显示 MM-DD"""
    try:
        import datetime
        dt = datetime.datetime.fromtimestamp(ts)
        now = datetime.datetime.now()
        today = now.date()
        d = dt.date()
        if d == today:
            return dt.strftime("%H:%M")
        if (today - d).days == 1:
            return "昨天"
        return dt.strftime("%m-%d")
    except Exception:
        return ""


class InputEdit(QTextEdit):
    """微信式输入框：Enter 发送，Shift+Enter 换行"""
    sendRequested = Signal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setObjectName("InputEdit")
        self.setAcceptRichText(False)
        self.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.setFixedHeight(42)
        self.document().documentLayout().documentSizeChanged.connect(self._auto_grow)

    def _auto_grow(self):
        h = int(self.document().size().height()) + 16
        self.setFixedHeight(min(max(h, 42), 132))

    def keyPressEvent(self, e):
        if e.key() in (Qt.Key_Return, Qt.Key_Enter) and not (e.modifiers() & Qt.ShiftModifier):
            self.sendRequested.emit()
            e.accept()
        else:
            super().keyPressEvent(e)


class SessionItemWidget(QWidget):
    """会话列表项：头像 + 标题 + 最后消息预览 + 时间（Fluent 卡片风，微信式布局）"""

    def __init__(self, title, preview, color, ts_text="", parent=None):
        super().__init__(parent)
        self._sel = False
        self.setObjectName("sessionItem")
        lay = QHBoxLayout(self)
        lay.setContentsMargins(10, 8, 10, 8)
        lay.setSpacing(10)
        first = (title or "新").strip()[:1].upper()
        lay.addWidget(AvatarLabel(first, color, size=38))
        col = QVBoxLayout()
        col.setSpacing(2)
        top = QHBoxLayout()
        top.setSpacing(6)
        t = QLabel(title)
        t.setObjectName("sessionTitle")
        t.setStyleSheet("font-size: 14px; font-weight: 600; color: #1f2937; background: transparent;")
        t.setWordWrap(False)
        top.addWidget(t, 1)
        ts = QLabel(ts_text)
        ts.setObjectName("sessionTime")
        ts.setStyleSheet("font-size: 11px; color: #9ca3af; background: transparent;")
        ts.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
        top.addWidget(ts)
        p = QLabel(preview)
        p.setObjectName("sessionPreview")
        p.setStyleSheet("font-size: 12px; color: #9ca3af; background: transparent;")
        p.setWordWrap(False)
        col.addLayout(top)
        col.addWidget(p)
        lay.addLayout(col, 1)

    def set_selected(self, sel):
        self._sel = sel
        self.setProperty("selected", "true" if sel else "false")
        st = self.style()
        st.unpolish(self)
        st.polish(self)


def pollinations_image(prompt, save_path, timeout=180):
    """Pollinations.ai 免 Key 文生图（URL 直调）"""
    import urllib.parse
    url = ("https://image.pollinations.ai/prompt/" + urllib.parse.quote(prompt)
           + "?width=1024&height=1024&model=flux&nologo=true")
    r = requests.get(url, timeout=timeout)
    if r.status_code == 200 and len(r.content) > 500:
        with open(save_path, "wb") as f:
            f.write(r.content)
        return save_path
    raise APIError(f"Pollinations 图片生成失败: HTTP {r.status_code}")


# ============ 消息气泡组件 ============
class VoiceBar(QFrame):
    """语音条：播放按钮 + 时长，点击播放/暂停"""

    def __init__(self, mp3_path, parent=None):
        super().__init__(parent)
        self.mp3_path = mp3_path
        self.setObjectName("voiceBar")
        self.setCursor(Qt.PointingHandCursor)
        lay = QHBoxLayout(self)
        lay.setContentsMargins(12, 6, 12, 6)
        lay.setSpacing(8)
        self.btn = QLabel("▶")
        self.btn.setStyleSheet("font-size:15px; color:#1a56db;")
        self.dur = QLabel("…")
        self.dur.setStyleSheet("color:#374151; font-size:12px;")
        lay.addWidget(self.btn)
        lay.addWidget(self.dur)
        lay.addStretch()
        self.player = QMediaPlayer(self)
        self.audio = QAudioOutput(self)
        self.player.setAudioOutput(self.audio)
        self.player.setSource(QUrl.fromLocalFile(mp3_path))

        def on_duration(d):
            if d > 0:
                sec = int(d / 1000)
                self.dur.setText(f"{sec//60}:{sec%60:02d}")

        def on_state(st):
            if st == QMediaPlayer.PlayingState:
                self.btn.setText("⏸")
            else:
                self.btn.setText("▶")

        self.player.durationChanged.connect(on_duration)
        self.player.playbackStateChanged.connect(on_state)

    def mousePressEvent(self, e):
        if self.player.playbackState() == QMediaPlayer.PlayingState:
            self.player.pause()
        else:
            self.player.play()
        super().mousePressEvent(e)

    def stop(self):
        self.player.stop()


class BubbleFrame(QFrame):
    """聊天气泡容器：文本 + 可追加语音条/搜索卡片/图片网格（Fluent 大圆角）"""

    MAX_BUBBLE_W = 720   # 气泡最大宽度（超过则自动换行）
    MIN_BUBBLE_W = 60    # 气泡最小宽度（短文本也不至于太窄）

    def __init__(self, who="ai", parent=None):
        super().__init__(parent)
        self.who = who
        self._row = None
        self.setObjectName("bubbleUser" if who == "user" else "bubbleAi")
        self.setMaximumWidth(self.MAX_BUBBLE_W)
        self.setSizePolicy(QSizePolicy.Maximum, QSizePolicy.Minimum)
        if FLUENT_OK:
            eff = QGraphicsDropShadowEffect(self)
            eff.setBlurRadius(24)
            eff.setOffset(0, 4)
            eff.setColor(QColor(31, 41, 55, 22) if who == "ai" else QColor(26, 86, 219, 40))
            self.setGraphicsEffect(eff)
        outer = QVBoxLayout(self)
        outer.setContentsMargins(16, 11, 16, 11)
        outer.setSpacing(6)
        self.inner = QVBoxLayout()
        self.inner.setContentsMargins(0, 0, 0, 0)
        self.inner.setSpacing(6)
        outer.addLayout(self.inner)
        self.text_view = QTextBrowser()
        self.text_view.setOpenExternalLinks(True)
        self.text_view.setWordWrapMode(QTextOption.WrapAtWordBoundaryOrAnywhere)
        _font = QFont()
        _font.setPixelSize(14)  # 与样式表 font-size:14px 一致，保证宽度测量准确
        self.text_view.setFont(_font)
        self.text_view.document().setDefaultStyleSheet("""
            p { margin: 4px 0; }
            code { background: #f0f0f0; padding: 1px 4px; border-radius: 3px; font-family: Consolas; }
            pre { background: #282c34; color: #abb2bf; padding: 10px; border-radius: 8px; }
        """)
        self.text_view.setFrameShape(QFrame.NoFrame)
        self.text_view.setStyleSheet("background: transparent; border: none; font-size: 14px;")
        self.inner.addWidget(self.text_view)

    def _fit_to_text(self):
        """气泡宽度自适应文本：短文本窄气泡，长文本撑到最大宽度后换行"""
        if not hasattr(self, "text_view"):
            return
        fm = self.text_view.fontMetrics()
        ideal = 0
        for ln in self.text_view.toPlainText().split("\n"):
            ideal = max(ideal, fm.horizontalAdvance(ln))
        # 富文本（加粗/代码块等）比纯文本略宽，用文档理想宽度兜底取更大值
        try:
            doc = self.text_view.document()
            dw = doc.idealWidth()
            if dw and dw > ideal:
                ideal = dw
        except Exception:
            pass
        if ideal <= 0:
            return
        # 内容 + 左右内边距(28) + 少量余量；上限封顶后由 QTextBrowser 自动换行
        w = int(ideal) + 28 + 8
        w = max(self.MIN_BUBBLE_W, min(self.MAX_BUBBLE_W, w))
        self.setFixedWidth(w)

    def set_text(self, text):
        self.text_view.setPlainText(text)
        self._fit_to_text()
        self._scroll_end()

    def append_text(self, text):
        self.text_view.insertPlainText(text)
        self.text_view.moveCursor(QTextCursor.End)
        self._fit_to_text()
        self._scroll_end()

    def _scroll_end(self):
        sb = self.text_view.verticalScrollBar()
        sb.setValue(sb.maximum())

    def add_voice_bar(self, mp3_path):
        if mp3_path and os.path.exists(mp3_path):
            self.setFixedWidth(self.MAX_BUBBLE_W)
            self.inner.addWidget(VoiceBar(mp3_path))

    def add_search_card(self, query, results):
        self.setFixedWidth(self.MAX_BUBBLE_W)
        card = QFrame()
        card.setObjectName("searchCard")
        lay = QVBoxLayout(card)
        lay.setContentsMargins(10, 8, 10, 8)
        lay.setSpacing(6)
        head = QLabel(f"🔍 搜索：{html_escape(query)}")
        head.setStyleSheet("color:#1a56db; font-weight:bold; font-size:13px;")
        lay.addWidget(head)
        if not results:
            empty = QLabel("没有搜到结果，建议换个关键词。")
            empty.setStyleSheet("color:#9ca3af;")
            lay.addWidget(empty)
        for it in results[:6]:
            t = QLabel(f'<a href="{html_escape(it.get("url", ""))}" style="color:#1a56db; text-decoration:none; font-weight:bold;">{html_escape(it.get("title", ""))}</a>')
            t.setOpenExternalLinks(True)
            t.setWordWrap(True)
            lay.addWidget(t)
            if it.get("snippet"):
                s = QLabel(html_escape(it.get("snippet", "")))
                s.setWordWrap(True)
                s.setStyleSheet("color:#6b7280; font-size:12px;")
                lay.addWidget(s)
        self.inner.addWidget(card)

    def add_image_grid(self, items):
        """items: [(local_path, thumb_url, orig_url, title)]"""
        if not items:
            return
        self.setFixedWidth(self.MAX_BUBBLE_W)
        grid = QGridLayout()
        grid.setSpacing(6)
        max_cols = 3
        for idx, (local, thumb, orig, title) in enumerate(items):
            thumb_label = QLabel()
            thumb_label.setFixedSize(150, 110)
            thumb_label.setAlignment(Qt.AlignCenter)
            pix = QPixmap(local)
            if not pix.isNull():
                thumb_label.setPixmap(pix.scaled(150, 110, Qt.KeepAspectRatio, Qt.SmoothTransformation))
            thumb_label.setStyleSheet("border:1px solid #e5e7eb; border-radius:6px; background:#f9fafb;")
            thumb_label.setCursor(Qt.PointingHandCursor)
            thumb_label.setToolTip(title)
            thumb_label.mousePressEvent = lambda e, u=orig: QDesktopServices.openUrl(QUrl(u))
            grid.addWidget(thumb_label, idx // max_cols, idx % max_cols)
        wrap = QWidget()
        wrap.setLayout(grid)
        self.inner.addWidget(wrap)

    def add_local_image(self, path):
        """展示本地图片（生成/分析预览）"""
        self.setFixedWidth(self.MAX_BUBBLE_W)
        pix = QPixmap(path)
        if pix.isNull():
            return
        img = QLabel()
        img.setAlignment(Qt.AlignCenter)
        img.setPixmap(pix.scaled(360, 360, Qt.KeepAspectRatio, Qt.SmoothTransformation))
        img.setCursor(Qt.PointingHandCursor)
        img.setToolTip("点击打开原图")
        img.mousePressEvent = lambda e, u=path: QDesktopServices.openUrl(QUrl.fromLocalFile(u))
        self.inner.addWidget(img)

    def add_video_card(self, path):
        """展示本地视频卡片（文生视频结果）：缩略图+文件名+播放/打开按钮"""
        self.setFixedWidth(self.MAX_BUBBLE_W)
        card = QFrame()
        card.setStyleSheet(
            "QFrame{background:#f9fafb;border:1px solid #e5e7eb;border-radius:10px;}"
            "QPushButton{background:#1a56db;color:white;border:none;border-radius:6px;padding:4px 14px;font-size:12px;}")
        lay = QVBoxLayout(card)
        lay.setContentsMargins(10, 10, 10, 10)
        lay.setSpacing(6)
        # 首帧缩略图（QMediaPlayer 可提取，但此处直接显示文件图标占位 + 文件名）
        info = QLabel(f"🎬 {os.path.basename(path)}\n{os.path.getsize(path)/1024/1024:.1f} MB")
        info.setStyleSheet("color:#374151; font-size:13px; background:transparent; border:none;")
        info.setWordWrap(True)
        lay.addWidget(info)
        row = QHBoxLayout()
        btn_play = QPushButton("▶ 播放")
        btn_play.setCursor(Qt.PointingHandCursor)
        btn_play.clicked.connect(lambda: QDesktopServices.openUrl(QUrl.fromLocalFile(path)))
        btn_open = QPushButton("📂 打开位置")
        btn_open.setStyleSheet(
            "QPushButton{background:#eef2ff;color:#1a56db;border:none;border-radius:6px;padding:4px 10px;font-size:12px;}")
        btn_open.setCursor(Qt.PointingHandCursor)
        btn_open.clicked.connect(lambda: QDesktopServices.openUrl(QUrl.fromLocalFile(os.path.dirname(path))))
        row.addWidget(btn_play)
        row.addWidget(btn_open)
        row.addStretch(1)
        lay.addLayout(row)
        self.inner.addWidget(card)

    def add_system_tip(self, text):
        tip = QLabel(text)
        tip.setStyleSheet("color:#6b7280; font-size:12px; font-style:italic;")
        self.inner.addWidget(tip)

    def add_regen_btn(self, callback):
        """添加「重新回答」小按钮（不满意可重新生成）"""
        self.setFixedWidth(self.MAX_BUBBLE_W)
        row = QHBoxLayout()
        row.setSpacing(4)
        btn = QPushButton("↻ 重新回答")
        btn.setCursor(Qt.PointingHandCursor)
        btn.setStyleSheet(
            "QPushButton{background:transparent;color:#9ca3af;border:1px solid #e5e7eb;"
            "border-radius:10px;padding:2px 10px;font-size:12px;}"
            "QPushButton:hover{color:#1a56db;border-color:#1a56db;background:#f0f5ff;}")
        btn.setFixedHeight(22)
        btn.clicked.connect(callback)
        row.addWidget(btn)
        row.addStretch(1)
        self.inner.addLayout(row)


# ============ 设置对话框 ============
class ProviderEditDialog(QDialog):
    def __init__(self, cfg=None, parent=None):
        super().__init__(parent)
        self.setWindowTitle("编辑 Provider")
        self.setMinimumWidth(520)
        cfg = cfg or {}
        self.cfg = dict(cfg)
        form = QFormLayout(self)
        self.p_name = QLineEdit(self.cfg.get("name", ""))
        self.p_type = QComboBox()
        self.p_type.addItems(["openai", "baidu", "g4f"])
        self.p_type.setCurrentText(self.cfg.get("type", "openai"))
        self.p_base = QLineEdit(self.cfg.get("base_url", ""))
        self.p_key = QLineEdit(self.cfg.get("api_key", ""))
        self.p_key.setEchoMode(QLineEdit.Password)
        self.p_secret = QLineEdit(self.cfg.get("secret_key", ""))
        self.p_secret.setEchoMode(QLineEdit.Password)
        self.p_chat = QLineEdit(self.cfg.get("chat_model", ""))
        self.p_img = QLineEdit(self.cfg.get("image_model", ""))
        form.addRow("名称", self.p_name)
        form.addRow("类型", self.p_type)
        form.addRow("Base URL", self.p_base)
        form.addRow("API Key", self.p_key)
        form.addRow("Secret Key（百度）", self.p_secret)
        form.addRow("对话模型", self.p_chat)
        form.addRow("图像模型（留空不支持）", self.p_img)
        btns = QHBoxLayout()
        ok = QPushButton("保存")
        cancel = QPushButton("取消")
        ok.clicked.connect(self._save)
        cancel.clicked.connect(self.reject)
        btns.addStretch()
        btns.addWidget(ok)
        btns.addWidget(cancel)
        form.addRow(btns)
        self.setStyleSheet(APP_STYLE)

    def _save(self):
        self.cfg.update({
            "name": self.p_name.text().strip(),
            "type": self.p_type.currentText(),
            "base_url": self.p_base.text().strip(),
            "api_key": self.p_key.text().strip(),
            "secret_key": self.p_secret.text().strip(),
            "chat_model": self.p_chat.text().strip(),
            "image_model": self.p_img.text().strip(),
        })
        self.accept()


class SettingsDialog(QDialog):
    def __init__(self, config, parent=None):
        super().__init__(parent)
        self.setWindowTitle("设置 - Provider 与语音")
        self.resize(760, 520)
        self.config = dict(config)
        self.providers = {}
        for pid, p in config.get("providers", {}).items():
            self.providers[pid] = dict(p)
        root = QVBoxLayout(self)

        tip = QLabel('备用通道免费注册：智谱 <a href="https://open.bigmodel.cn">open.bigmodel.cn</a>（GLM-4-Flash 免费）｜百度 <a href="https://console.bce.baidu.com/qianfan">千帆</a>（ERNIE-Speed 免费）｜本地 <a href="https://ollama.com">Ollama</a>（完全离线）。主 API 故障时自动切换。')
        tip.setOpenExternalLinks(True)
        tip.setWordWrap(True)
        tip.setStyleSheet("color:#6b7280; font-size:12px;")
        root.addWidget(tip)

        self.table = QTableWidget(0, 5)
        self.table.setHorizontalHeaderLabels(["启用", "名称", "类型", "Base URL", "对话模型"])
        self.table.horizontalHeader().setSectionResizeMode(3, QHeaderView.Stretch)
        self.table.verticalHeader().setVisible(False)
        self.table.setSelectionBehavior(QTableWidget.SelectRows)
        root.addWidget(self.table, 1)
        self.reload_table()

        btn_row = QHBoxLayout()
        add_btn = QPushButton("添加 Provider")
        edit_btn = QPushButton("编辑选中")
        del_btn = QPushButton("删除选中")
        add_btn.clicked.connect(self._add)
        edit_btn.clicked.connect(self._edit)
        del_btn.clicked.connect(self._del)
        btn_row.addWidget(add_btn)
        btn_row.addWidget(edit_btn)
        btn_row.addWidget(del_btn)
        btn_row.addStretch()
        root.addLayout(btn_row)

        self.tts_check = QCheckBox("启用语音回复（语音条，edge-tts 免费）")
        self.tts_check.setChecked(self.config.get("tts_enabled", True))
        root.addWidget(self.tts_check)

        btns = QHBoxLayout()
        ok = QPushButton("保存")
        cancel = QPushButton("取消")
        ok.clicked.connect(self._save)
        cancel.clicked.connect(self.reject)
        btns.addStretch()
        btns.addWidget(ok)
        btns.addWidget(cancel)
        root.addLayout(btns)
        self.setStyleSheet(APP_STYLE)

    def reload_table(self):
        self.table.setRowCount(0)
        for pid, p in self.providers.items():
            row = self.table.rowCount()
            self.table.insertRow(row)
            chk = QCheckBox()
            chk.setChecked(bool(p.get("enabled", True)))
            chk.setStyleSheet("margin-left:8px;")
            chk.stateChanged.connect(lambda st, r=row: self._toggle_enabled(r, st))
            self.table.setCellWidget(row, 0, chk)
            self.table.setItem(row, 1, QTableWidgetItem(p.get("name", pid)))
            self.table.setItem(row, 2, QTableWidgetItem(p.get("type", "openai")))
            self.table.setItem(row, 3, QTableWidgetItem(p.get("base_url", "")))
            self.table.setItem(row, 4, QTableWidgetItem(p.get("chat_model", "")))
            self.table.item(row, 1).setData(Qt.UserRole, pid)

    def _toggle_enabled(self, row, state):
        pid = self.table.item(row, 1).data(Qt.UserRole)
        if pid in self.providers:
            self.providers[pid]["enabled"] = bool(state)

    def _current_pid(self):
        row = self.table.currentRow()
        if row < 0:
            QMessageBox.information(self, "提示", "请先选择一行")
            return None
        return self.table.item(row, 1).data(Qt.UserRole)

    def _add(self):
        dlg = ProviderEditDialog(parent=self)
        if dlg.exec():
            pid = f"custom_{int(time.time())}"
            self.providers[pid] = dlg.cfg
            self.reload_table()

    def _edit(self):
        pid = self._current_pid()
        if not pid:
            return
        dlg = ProviderEditDialog(self.providers[pid], self)
        if dlg.exec():
            self.providers[pid] = dlg.cfg
            self.reload_table()

    def _del(self):
        pid = self._current_pid()
        if not pid:
            return
        if pid == "agnes":
            QMessageBox.information(self, "提示", "主 Provider 不可删除，可取消启用")
            return
        del self.providers[pid]
        self.reload_table()

    def _save(self):
        self.config["providers"] = self.providers
        self.config["tts_enabled"] = self.tts_check.isChecked()
        self.accept()


# ============ 主窗口 ============
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.config = load_config()
        self.mode = self.config.get("mode", "cloud")
        self.clients = build_clients(self.config, self.mode)
        self.messages = []
        self.busy = False
        self.current_image_path = None
        self.last_full_text = ""
        self.last_ai_bubble = None
        self.streaming_active = False
        self.pending_files = []
        self.attach_labels = []
        self.session_files = set()  # 当前会话产生的缓存文件（图片/视频/语音/搜索缓存），删除会话时一并清理
        self.dept_state = {}
        for dept in DEPARTMENTS:
            self.dept_state[dept] = self.config.get("dept_disabled", {}).get(dept, True)
        self.setAcceptDrops(True)  # 支持把文件拖进窗口
        self.init_ui()
        self._update_mode_btn()
        self.load_history_list()

    def init_ui(self):
        self.setWindowTitle("AgnesAgent - 免费全能 AI 桌面助手")
        self.resize(1180, 780)
        self.setMinimumSize(860, 600)
        self.setStyleSheet(APP_STYLE)

        central = QWidget()
        self.setCentralWidget(central)
        root = QVBoxLayout(central)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        # 顶栏（Fluent 风格）
        topbar = QFrame()
        topbar.setObjectName("topbar")
        tl = QHBoxLayout(topbar)
        tl.setContentsMargins(20, 10, 20, 10)
        tl.setSpacing(10)
        self.avatar_bar = AvatarLabel("A", "#2b5bff", size=34)
        tl.addWidget(self.avatar_bar)
        title = QLabel("AgnesAgent")
        title.setObjectName("appTitle")
        tl.addWidget(title)
        dot = QLabel("●")
        dot.setObjectName("onlineDot")
        tl.addWidget(dot)
        self.mode_btn = QPushButton()
        self.mode_btn.setObjectName("modeBtn")
        self.mode_btn.setCheckable(True)
        self.mode_btn.setCursor(Qt.PointingHandCursor)
        self.mode_btn.clicked.connect(self.toggle_mode)
        tl.addWidget(self.mode_btn)
        tl.addStretch()
        self.status_label = QLabel("就绪")
        self.status_label.setObjectName("statusLabel")
        tl.addWidget(self.status_label)
        btn_settings = QPushButton("设置")
        btn_settings.setObjectName("ghostBtn")
        btn_settings.setCursor(Qt.PointingHandCursor)
        btn_settings.clicked.connect(self.open_settings)
        tl.addWidget(btn_settings)
        root.addWidget(topbar)

        # 主体
        splitter = QSplitter(Qt.Horizontal)
        splitter.setHandleWidth(1)
        root.addWidget(splitter, 1)

        # 侧边栏（微信式会话列表 + Fluent 部门挂载区）
        sidebar = QFrame()
        sidebar.setObjectName("sidebar")
        sidebar.setFixedWidth(272)
        sl = QVBoxLayout(sidebar)
        sl.setContentsMargins(12, 14, 12, 10)
        sl.setSpacing(8)
        new_btn = QPushButton("＋ 新建会话")
        new_btn.setObjectName("primaryBtn")
        new_btn.setCursor(Qt.PointingHandCursor)
        new_btn.clicked.connect(self.new_session)
        sl.addWidget(new_btn)
        self.history_list = QListWidget()
        self.history_list.setObjectName("sessionList")
        self.history_list.itemClicked.connect(self.load_session)
        self.history_list.currentRowChanged.connect(self._on_session_selected)
        self.history_list.setContextMenuPolicy(Qt.CustomContextMenu)
        self.history_list.customContextMenuRequested.connect(self._history_menu)
        self.history_list.setSpacing(3)
        sl.addWidget(self.history_list, 1)

        # 部门挂载区：每个功能 = 一个部门，Fluent 开关卡片，部门内嵌快捷动作按钮
        dept_scroll = QScrollArea()
        dept_scroll.setWidgetResizable(True)
        dept_scroll.setFrameShape(QFrame.NoFrame)
        dept_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")
        dept_frame = QWidget()
        dept_frame.setObjectName("deptFrame")
        dl = QVBoxLayout(dept_frame)
        dl.setContentsMargins(10, 8, 10, 8)
        dl.setSpacing(5)
        dept_title = QLabel("⚡ 功能部门")
        dept_title.setObjectName("deptTitle")
        dl.addWidget(dept_title)
        self.dept_checkboxes = {}
        self._action_btns = {}
        for dept, info in DEPARTMENTS.items():
            block = QFrame()
            block.setObjectName("deptBlock")
            bl = QVBoxLayout(block)
            bl.setContentsMargins(12, 7, 12, 7)
            bl.setSpacing(4)
            row = QHBoxLayout()
            row.setSpacing(6)
            name_lbl = QLabel(f"{info['icon']} {dept}")
            name_lbl.setObjectName("deptName")
            name_lbl.setToolTip(info.get("desc", ""))
            row.addWidget(name_lbl)
            row.addStretch()
            if FLUENT_OK:
                sw = SwitchButton()
                sw.setChecked(bool(self.dept_state.get(dept, True)))
                sw.checkedChanged.connect(lambda checked, d=dept: self._toggle_dept(d, checked))
                sw.setToolTip(info.get("desc", ""))
                row.addWidget(sw)
                cb = sw
            else:
                cb = QCheckBox(f"{info['icon']} {dept}")
                cb.setToolTip(info.get("desc", ""))
                cb.setChecked(bool(self.dept_state.get(dept, True)))
                cb.toggled.connect(lambda checked, d=dept: self._toggle_dept(d, checked))
                row.addWidget(cb)
            bl.addLayout(row)
            for aid in info.get("actions", []):
                label, handler = ACTION_MAP.get(aid, (aid, None))
                if handler:
                    ab = QPushButton(label)
                    ab.setObjectName("deptAction")
                    ab.setCursor(Qt.PointingHandCursor)
                    ab.clicked.connect(lambda _, h=handler: h())
                    bl.addWidget(ab)
                    self._action_btns[aid] = ab
            self.dept_checkboxes[dept] = cb
            dl.addWidget(block)
        dl.addStretch()
        dept_scroll.setWidget(dept_frame)
        sl.addWidget(dept_scroll, 2)
        # 按当前部门开关初始化动作按钮启用状态
        enabled_actions = get_enabled_actions(self.dept_state)
        for aid, ab in self._action_btns.items():
            ab.setEnabled(aid in enabled_actions)

        hint = QLabel("右键会话可删除/重命名")
        hint.setObjectName("hint")
        sl.addWidget(hint)
        splitter.addWidget(sidebar)

        # 聊天区（微信式浅灰背景）
        chat_area = QWidget()
        chat_area.setObjectName("chatArea")
        cl = QVBoxLayout(chat_area)
        cl.setContentsMargins(0, 0, 0, 0)
        cl.setSpacing(0)
        self.scroll = QScrollArea()
        self.scroll.setWidgetResizable(True)
        self.scroll.setFrameShape(QFrame.NoFrame)
        self.scroll.setStyleSheet("QScrollArea { background: #f5f7fa; border: none; }")
        self.chat_container = QWidget()
        self.chat_container.setObjectName("chatContainer")
        self.chat_layout = QVBoxLayout(self.chat_container)
        self.chat_layout.setContentsMargins(28, 18, 28, 18)
        self.chat_layout.setSpacing(14)
        self.chat_layout.addStretch(1)
        self.scroll.setWidget(self.chat_container)
        cl.addWidget(self.scroll, 1)

        # 输入区（Fluent 圆角输入容器）
        input_frame = QFrame()
        input_frame.setObjectName("inputFrame")
        il = QVBoxLayout(input_frame)
        il.setContentsMargins(20, 10, 20, 14)
        il.setSpacing(8)

        quick_row = QHBoxLayout()
        quick_row.setSpacing(8)
        btn_attach = QPushButton("📎 文件")
        btn_attach.setObjectName("chip")
        btn_attach.setCursor(Qt.PointingHandCursor)
        btn_attach.setToolTip("选择本地文件（txt/pdf/docx/xlsx/csv/md/py/图片等），输入问题后 AI 自动读取处理")
        btn_attach.clicked.connect(self.attach_files)
        quick_row.addWidget(btn_attach)
        quick_row.addStretch()
        il.addLayout(quick_row)

        # 附件 chip 行（拖拽/选择文件后显示，可一键清空）
        self.attach_container = QWidget()
        self.attach_container.setVisible(False)
        ac = QHBoxLayout(self.attach_container)
        ac.setContentsMargins(0, 0, 0, 0)
        ac.setSpacing(6)
        self.attach_container_layout = ac
        self.attach_clear_btn = QPushButton("✕ 清空附件")
        self.attach_clear_btn.setObjectName("chip")
        self.attach_clear_btn.setToolTip("清空已选附件")
        self.attach_clear_btn.clicked.connect(self.clear_pending_files)
        self.attach_labels = []
        il.addWidget(self.attach_container)

        row = QHBoxLayout()
        row.setSpacing(10)
        self.input_box = InputEdit()
        self.input_box.setPlaceholderText("输入你的问题，Enter 发送，Shift+Enter 换行；支持搜索、生成图片、分析图片、语音输入…")
        self.input_box.sendRequested.connect(self.send_message)
        self.send_btn = QPushButton("发送")
        self.send_btn.setObjectName("primaryBtn")
        self.send_btn.setCursor(Qt.PointingHandCursor)
        self.send_btn.clicked.connect(self.send_message)
        row.addWidget(self.input_box, 1)
        row.addWidget(self.send_btn)
        il.addLayout(row)
        cl.addWidget(input_frame)

        splitter.addWidget(chat_area)
        splitter.setSizes([272, 908])


    # ---------- 会话管理 ----------
    def session_file(self, sid):
        return os.path.join(HISTORY_DIR, f"{sid}.json")

    def load_history_list(self):
        self.history_list.blockSignals(True)
        self.history_list.clear()
        colors = ["#2b5bff", "#7c3aed", "#059669", "#d97706", "#dc2626", "#0891b2", "#db2777"]
        if os.path.exists(HISTORY_DIR):
            # 按最后活动时间（文件修改时间）排序，最新会话在最上，符合主流聊天软件逻辑
            files = []
            for name in os.listdir(HISTORY_DIR):
                if name.endswith(".json"):
                    p = os.path.join(HISTORY_DIR, name)
                    try:
                        files.append((os.path.getmtime(p), p, name))
                    except Exception:
                        pass
            files.sort(key=lambda x: x[0], reverse=True)
            for i, (mtime, path, name) in enumerate(files):
                try:
                    with open(path, "r", encoding="utf-8") as f:
                        data = json.load(f)
                    sid = data.get("id", name[:-5])
                    title = data.get("title", "新会话")
                    msgs = data.get("messages", [])
                    last = msgs[-1].get("content", "") if msgs else ""
                    if isinstance(last, list):
                        last = "".join(p.get("text", "") for p in last if isinstance(p, dict))
                    preview = last[:60].replace("\n", " ") if last else ""
                    item = QListWidgetItem()
                    item.setSizeHint(QSize(0, 64))
                    item.setData(Qt.UserRole, sid)
                    color = colors[i % len(colors)]
                    self.history_list.addItem(item)
                    self.history_list.setItemWidget(
                        item, SessionItemWidget(title, preview, color, ts_text=_fmt_session_time(mtime)))
                except Exception:
                    pass
        self.history_list.blockSignals(False)

    def _history_menu(self, pos):
        item = self.history_list.itemAt(pos)
        if not item:
            return
        menu = QMenu(self)
        act_del = menu.addAction("删除此会话")
        act_rename = menu.addAction("重命名会话")
        act = menu.exec(self.history_list.mapToGlobal(pos))
        sid = item.data(Qt.UserRole)
        if act == act_del:
            self._delete_session_with_cache(sid)
        elif act == act_rename:
            from PySide6.QtWidgets import QInputDialog
            new_title, ok = QInputDialog.getText(self, "重命名", "新标题：", text=item.text())
            if ok and new_title.strip():
                try:
                    path = self.session_file(sid)
                    with open(path, "r", encoding="utf-8") as f:
                        data = json.load(f)
                    data["title"] = new_title.strip()
                    with open(path, "w", encoding="utf-8") as f:
                        json.dump(data, f, ensure_ascii=False, indent=2)
                    self.load_history_list()
                except Exception:
                    pass

    def _delete_session_with_cache(self, sid):
        """删除会话并连带清理该会话产生的所有缓存文件（图片/视频/语音/搜索缓存）"""
        path = self.session_file(sid)
        cache_files = set()
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            cache_files.update(data.get("files", []) or [])
            # 兼容旧会话：从消息文本中提取出现在缓存目录下的路径
            cache_dirs = (IMAGE_DIR, VIDEO_DIR, TTS_CACHE_DIR, SEARCH_CACHE_DIR)
            for m in data.get("messages", []):
                content = m.get("content", "")
                if isinstance(content, list):
                    for part in content:
                        if isinstance(part, dict) and part.get("type") == "image" and part.get("path"):
                            p = part["path"]
                            if any(p.startswith(d) for d in cache_dirs):
                                cache_files.add(p)
                elif isinstance(content, str):
                    for p in re.findall(r"[A-Za-z]:\\[^\s\"']+", content):
                        p = p.rstrip("，。；、,.;:：")
                        if any(p.startswith(d) for d in cache_dirs) and os.path.isfile(p):
                            cache_files.add(p)
        except Exception:
            pass
        # 只删除缓存目录下的文件，避免误删用户自己的文件
        for cf in cache_files:
            try:
                if cf and os.path.isfile(cf):
                    os.remove(cf)
            except Exception:
                pass
        try:
            os.remove(path)
        except Exception:
            pass
        if getattr(self, "session_id", None) == sid:
            self.messages = []
            self.session_files = set()
            self.session_id = None
            self.clear_chat()
            self.new_session()
        self.load_history_list()

    def new_session(self):
        self.messages = []
        self._media_pending = []
        self.current_image_path = None
        self.clear_chat()
        self.session_id = f"chat_{time.strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:6]}"
        self.session_files = set()
        self.last_full_text = ""
        self.show_ai("你好，我是 AgnesAgent，由 MIRAGE 独立开发的免费全能 AI 桌面助手。我可以对话问答、联网搜索（网页+图片）、执行代码、读取本地文件、生成图片、分析图片，全程免费。有什么可以帮你？")

    def clear_chat(self):
        while self.chat_layout.count() > 1:
            item = self.chat_layout.takeAt(0)
            w = item.widget()
            if w:
                w.deleteLater()
        self.last_ai_bubble = None
        self.streaming_active = False

    def load_session(self, item):
        sid = item.data(Qt.UserRole)
        path = self.session_file(sid)
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            self.messages = data.get("messages", [])
            self.session_id = sid
            self.session_files = set(data.get("files", []) or [])
            self._media_pending = []
            self.clear_chat()
            img_exts = (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp")
            vid_exts = (".mp4", ".avi", ".mkv", ".mov", ".webm")
            for m in self.messages:
                if m.get("role") == "user":
                    content = m.get("content", "")
                    if isinstance(content, list):
                        content = content[0].get("text", "")
                    b = self.show_user(content)
                    # 渲染用户消息中的本地图片附件（按消息顺序，不堆到末尾）
                    if isinstance(content, str):
                        for p in re.findall(r"路径\s*([A-Za-z]:\\[^\s\"'】]+)", content):
                            p = p.rstrip("，。；、,.;:：")
                            if os.path.isfile(p) and os.path.splitext(p)[1].lower() in img_exts:
                                b.add_local_image(p)
                elif m.get("role") == "assistant" and m.get("content"):
                    self.show_ai(m["content"])
                    # 按消息顺序渲染该条回复关联的媒体（生成图片/视频/语音）
                    for med in (m.get("media") or []):
                        p = med.get("path", "") if isinstance(med, dict) else med
                        t = med.get("type", "") if isinstance(med, dict) else ""
                        if not p:
                            continue
                        if not os.path.isfile(p):
                            cand = os.path.join(IMAGE_DIR if t == "image" else VIDEO_DIR, os.path.basename(p))
                            if os.path.isfile(cand):
                                p = cand
                            else:
                                continue
                        ext = os.path.splitext(p)[1].lower()
                        if t == "image" or ext in img_exts:
                            self._add_bubble("ai").add_local_image(p)
                        elif t == "video" or ext in vid_exts:
                            self._add_bubble("ai").add_video_card(p)
                        elif t == "audio" or ext in (".mp3", ".wav"):
                            self._add_bubble("ai").add_voice_bar(p)
            # 兼容旧会话：消息中没有 media 索引时，才把 files 里剩余媒体统一追加到末尾
            if not any(m.get("media") for m in self.messages):
                for p in sorted(self.session_files):
                    if p.startswith(SEARCH_CACHE_DIR) or "tts_cache" in p.replace("\\", "/"):
                        continue
                    ext = os.path.splitext(p)[1].lower()
                    is_img = ext in img_exts
                    is_vid = ext in vid_exts
                    if not is_img and not is_vid:
                        continue
                    if not os.path.isfile(p):
                        # 兼容 exe 迁移后路径失效：按文件名在当前 images/videos 目录查找
                        cand = os.path.join(IMAGE_DIR if is_img else VIDEO_DIR, os.path.basename(p))
                        if os.path.isfile(cand):
                            p = cand
                        else:
                            continue
                    if is_img:
                        self._add_bubble("ai").add_local_image(p)
                    elif is_vid:
                        self._add_bubble("ai").add_video_card(p)
        except Exception as e:
            QMessageBox.warning(self, "提示", f"加载会话失败: {e}")

    def save_current_session(self, title=None):
        if not getattr(self, "session_id", None):
            return
        path = self.session_file(self.session_id)
        try:
            data = {"id": self.session_id,
                    "title": title or (self.messages[-1].get("content", "")[:20] if self.messages else "新会话"),
                    "messages": self.messages,
                    "files": sorted(self.session_files)}
            with open(path, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    # ---------- UI 输出 ----------
    def _add_bubble(self, who="ai"):
        row = MessageRow(who)
        row.setMaximumWidth(780)
        row.set_time(current_time_str())
        row.bubble._row = row
        self.chat_layout.insertWidget(self.chat_layout.count() - 1, row)
        self._scroll_to_bottom()
        return row.bubble

    def _scroll_to_bottom(self):
        sb = self.scroll.verticalScrollBar()
        sb.setValue(sb.maximum())

    def show_ai(self, text):
        b = self._add_bubble("ai")
        b.set_text(text)
        self.last_ai_bubble = b
        return b

    def show_user(self, text):
        return self._add_bubble("user").set_text(text)

    def show_system_tip(self, text):
        b = self._add_bubble("ai")
        b.add_system_tip(text)

    def _on_session_selected(self, row):
        """会话选中时自动加载（配合微信式列表高亮）"""
        item = self.history_list.item(row)
        if item is not None:
            self.load_session(item)
        # 刷新选中高亮：仅当前会话项高亮
        for r in range(self.history_list.count()):
            it = self.history_list.item(r)
            w = self.history_list.itemWidget(it)
            if isinstance(w, SessionItemWidget):
                w.set_selected(r == row)

    def set_status(self, text):
        self.status_label.setText(text)

    def _update_mode_btn(self):
        """刷新模式开关按钮显示（物理隔离：开=云端在线，关=本地离线）"""
        if self.mode == "local":
            self.mode_btn.setText("🖥 本地模式（离线）")
            self.mode_btn.setChecked(False)
        else:
            self.mode_btn.setText("☁ 云端模式（在线）")
            self.mode_btn.setChecked(True)

    def toggle_mode(self):
        """本地/云端物理隔离开关：开启=云端，关闭=本地。
        切换时自动保存并清空当前会话上下文，新建会话——本地与在线严格二选一，
        绝不让上一模式的历史消息（尤其是在线 API 话术）污染另一模式的模型。"""
        new_mode = "local" if self.mode != "local" else "cloud"
        # 先保存当前会话（含媒体索引），再隔离清空
        try:
            self.save_current_session()
        except Exception:
            pass
        self.mode = new_mode
        self.config["mode"] = self.mode
        save_config(self.config)
        self.clients = build_clients(self.config, self.mode)
        # 清空当前上下文并新建会话，避免两模式聊天框并存/互相污染
        self.messages = []
        self.session_files = set()
        self.new_session()
        # 本地优先时自动启用本地大模型部门
        if self.mode == "local" and "本地大模型部" in self.dept_state:
            self.dept_state["本地大模型部"] = True
            self.config.setdefault("dept_disabled", {})["本地大模型部"] = True
            cb = self.dept_checkboxes.get("本地大模型部")
            if cb is not None:
                cb.setChecked(True)
            save_config(self.config)
        self._update_mode_btn()
        self.set_status("已切换到" + ("本地模式（离线，绝不联网）" if self.mode == "local" else "云端模式（在线）"))

    # ---------- 动作 ----------
    def _toggle_dept(self, dept, checked):
        """部门开关：关闭的部门其工具不再传给大模型，对应快捷按钮禁用"""
        self.dept_state[dept] = checked
        self.config.setdefault("dept_disabled", {})[dept] = checked
        save_config(self.config)
        if getattr(self, "_action_btns", None):
            enabled = get_enabled_actions(self.dept_state)
            for aid, btn in self._action_btns.items():
                btn.setEnabled(aid in enabled)

    def quick_action(self, prefix):
        self.input_box.setPlainText(prefix)
        self.input_box.setFocus()
        cursor = self.input_box.textCursor()
        cursor.movePosition(QTextCursor.End)
        self.input_box.setTextCursor(cursor)

    def ask_image_prompt(self):
        from PySide6.QtWidgets import QInputDialog
        text, ok = QInputDialog.getText(self, "生成图片（云端）", "请输入图片描述（支持中文）：")
        if ok and text.strip():
            self.generate_image(text.strip())

    def ask_image_prompt_local(self):
        from PySide6.QtWidgets import QInputDialog
        text, ok = QInputDialog.getText(self, "生成图片（本地）", "请输入图片描述（纯 CPU 离线生成，约 1-2 分钟）：")
        if ok and text.strip():
            self.generate_image_local(text.strip())

    def generate_image(self, prompt):
        if self.busy:
            return
        client = pick_image_client(self.clients)
        if not client:
            QMessageBox.warning(self, "提示", "当前没有可用的图像生成 Provider，请在设置中为某个 Provider 填写图像模型。")
            return
        self.busy = True
        self.set_status("正在生成图片…")
        self.show_user(f"🎨 {prompt}")
        self.worker = ImageWorker(client, prompt, IMAGE_DIR, self)
        self.worker.done.connect(self.on_image_done)
        self.worker.error.connect(self.on_image_error)
        self.worker.start()

    def generate_image_local(self, prompt):
        """本地生图：sd.cpp + Realistic Vision V6（纯 CPU 离线）"""
        if self.busy:
            return
        info, err = _locate_sd_engine()
        if err:
            QMessageBox.warning(self, "提示", err)
            return
        self.busy = True
        self.set_status("正在本地生成图片…")
        self.show_user(f"🖌 {prompt}")
        self.worker = LocalImageWorker(prompt, IMAGE_DIR, self)
        self.worker.done.connect(self.on_image_done)
        self.worker.error.connect(self.on_image_error)
        self.worker.start()
    def on_image_done(self, path):
        self.set_status("图片生成完成")
        b = self.show_ai(f"图片已生成：{os.path.basename(path)}")
        b.add_local_image(path)
        self.messages.append({"role": "user", "content": f"[生成图片] {path}"})
        self.messages.append({"role": "assistant", "content": f"图片已生成：{path}"})
        self.session_files.add(path)
        self.save_current_session()
        self.load_history_list()
        self.busy = False

    def on_image_error(self, msg):
        self.set_status("图片生成失败")
        self.show_ai(f"图片生成失败：{msg}")
        self.busy = False

    def ask_video_prompt(self):
        from PySide6.QtWidgets import QInputDialog
        text, ok = QInputDialog.getText(self, "生成视频（云端）", "请输入视频描述（主体、动作、场景）：")
        if ok and text.strip():
            self.generate_video(text.strip())

    def ask_video_prompt_local(self):
        from PySide6.QtWidgets import QInputDialog
        text, ok = QInputDialog.getText(self, "生成视频（本地）", "请输入视频描述（本地 ComfyUI + Wan2.1 自动调度，纯 CPU 约 20-40 分钟）：")
        if ok and text.strip():
            self.generate_video_local(text.strip())

    def generate_video(self, prompt):
        """文生视频（云端，失败自动切本地）：启动 VideoWorker，等待计时由 UI 每 10s 刷新"""
        if self.busy:
            return
        self.busy = True
        self.set_status("正在生成视频…")
        self.show_user(f"🎬 {prompt}")
        self._video_wait_start = time.time()
        self._video_wait_timer = QTimer(self)
        self._video_wait_timer.timeout.connect(self._update_video_wait)
        self._video_wait_timer.start(10000)
        b = self.show_ai("正在生成视频，请稍候…（等待计时中）")
        self._video_wait_bubble = b
        self._video_wait_text = "正在连接 AgnesAI 视频服务…"
        self.worker = VideoWorker(prompt, self)
        self.worker.progress.connect(self.on_video_progress)
        self.worker.done.connect(self.on_video_done)
        self.worker.error.connect(self.on_video_error)
        self.worker.start()

    def generate_video_local(self, prompt):
        """本地生视频：ComfyUI + Wan2.1 自动定位/启动/提交/下载，带真实进度"""
        if self.busy:
            return
        self.busy = True
        self.set_status("正在生成本地视频…")
        self.show_user(f"🎞️ {prompt}")
        self._video_wait_start = time.time()
        self._video_wait_timer = QTimer(self)
        self._video_wait_timer.timeout.connect(self._update_video_wait)
        self._video_wait_timer.start(10000)
        b = self.show_ai("正在调度本地 ComfyUI…（等待计时中）")
        self._video_wait_bubble = b
        self._video_wait_text = "正在定位 ComfyUI…"
        self.worker = LocalVideoWorker(prompt, 5, self)
        self.worker.progress.connect(self.on_video_progress)
        self.worker.done.connect(self.on_video_done)
        self.worker.error.connect(self.on_video_error)
        self.worker.start()

    def open_model_market(self):
        """模型市场：一键下载 Wan2.1 系列模型到本地 ComfyUI"""
        dlg = ModelMarketDialog(self)
        dlg.exec()

    def _update_video_wait(self):
        """每 10s 刷新一次等待计时（0:00 格式），只要没产出就一直显示正在产出"""
        if getattr(self, "_video_wait_bubble", None) is None:
            return
        el = int(time.time() - self._video_wait_start)
        mm, ss = divmod(el, 60)
        self._video_wait_bubble.set_text(f"{self._video_wait_text}\n已等待 {mm}:{ss:02d}，正在产出视频，请勿关闭窗口…")

    def on_video_progress(self, msg):
        self._video_wait_text = msg
        self._update_video_wait()

    def on_video_done(self, path, msg):
        if getattr(self, "_video_wait_timer", None):
            self._video_wait_timer.stop()
        self._video_wait_bubble = None
        self.set_status("视频生成完成")
        b = self.show_ai(msg)
        b.add_video_card(path)
        self.messages.append({"role": "user", "content": f"[生成视频] {path}"})
        self.messages.append({"role": "assistant", "content": f"{msg}：{path}"})
        self.session_files.add(path)
        self.save_current_session()
        self.load_history_list()
        self.busy = False

    def on_video_error(self, msg):
        if getattr(self, "_video_wait_timer", None):
            self._video_wait_timer.stop()
        self._video_wait_bubble = None
        self.set_status("视频生成失败")
        self.show_ai(f"视频生成失败：{msg}")
        self.busy = False

    def analyze_image(self):
        if self.busy:
            return
        path, _ = QFileDialog.getOpenFileName(self, "选择图片", "", "图片文件 (*.png *.jpg *.jpeg *.bmp *.webp *.gif)")
        if not path:
            return
        self.busy = True
        self.set_status("正在分析图片…")
        self.show_user(f"📷 分析图片：{os.path.basename(path)}")
        b = self.show_ai("正在识别图片内容…")
        b.add_local_image(path)
        self.worker = VisionWorker(self.clients[0], path, self)
        self.worker.done.connect(self.on_vision_done)
        self.worker.error.connect(self.on_vision_error)
        self.worker.start()

    def on_vision_done(self, text):
        self.set_status("分析完成")
        if self.last_ai_bubble:
            self.last_ai_bubble.set_text(text)
        else:
            self.show_ai(text)
        self.messages.append({"role": "user", "content": "[用户上传了一张图片]"})
        self.messages.append({"role": "assistant", "content": text})
        self.save_current_session()
        self.load_history_list()
        self.busy = False

    def on_vision_error(self, msg):
        self.set_status("分析失败")
        self.show_ai(f"图片分析失败：{msg}")
        self.busy = False

    def send_message(self):
        text = self.input_box.toPlainText().strip()
        if not text and not getattr(self, "pending_files", None):
            return
        if self.busy:
            return
        self.input_box.clear()
        self.run_chat(text)

    def add_pending_files(self, paths):
        if not getattr(self, "pending_files", None):
            self.pending_files = []
        self.pending_files.extend(paths)
        self._refresh_attach_chips()
        names = "、".join(os.path.basename(p) for p in paths)
        self.show_system_tip(f"📎 已添加附件：{names}。输入问题后发送，AI 会自动读取处理。")

    def clear_pending_files(self):
        self.pending_files = []
        self._refresh_attach_chips()

    def _refresh_attach_chips(self):
        for lb in getattr(self, "attach_labels", []):
            self.attach_container_layout.removeWidget(lb)
            lb.deleteLater()
        self.attach_labels = []
        if not getattr(self, "pending_files", None) or not self.pending_files:
            self.attach_container.setVisible(False)
            return
        for p in self.pending_files:
            lb = QLabel(f"📎 {os.path.basename(p)}")
            lb.setStyleSheet(
                "background:#e8f0fe;color:#1a56db;border-radius:8px;padding:2px 8px;font-size:12px;")
            self.attach_container_layout.addWidget(lb)
            self.attach_labels.append(lb)
        self.attach_container_layout.addWidget(self.attach_clear_btn)
        self.attach_container.setVisible(True)

    def attach_files(self):
        """选择文件挂载到待发送队列（含 chip 预览）"""
        if self.busy:
            return
        paths, _ = QFileDialog.getOpenFileNames(
            self, "选择文件（txt/pdf/docx/xlsx/csv/md/py/json/图片等）",
            "", "所有文件 (*.*)")
        if paths:
            self.add_pending_files(paths)

    def _read_attach(self, path):
        """读取附件：返回 (内容, 类型)；类型 text / image / error / other"""
        ext = os.path.splitext(path)[1].lower()
        img_exts = {".png", ".jpg", ".jpeg", ".bmp", ".webp", ".gif"}
        if ext in img_exts:
            return "", "image"
        text_exts = {".txt", ".md", ".py", ".json", ".csv", ".xml", ".log", ".html",
                     ".yaml", ".yml", ".ini", ".cfg", ".pdf", ".docx", ".xlsx", ".xlsm"}
        if ext not in text_exts:
            return "", "other"
        try:
            content = read_local_file(path)
            if content.startswith(("文件不存在", "文件过大", "读取失败", "路径不存在")):
                return content, "error"
            return content, "text"
        except Exception as e:
            return f"读取失败: {e}", "error"

    @staticmethod
    def _image_block(path):
        """将图片转为 base64 data URI，供多模态视觉消息使用；超过 5MB 跳过以防阻塞"""
        try:
            size = os.path.getsize(path)
            if size > 5 * 1024 * 1024:
                return {"type": "text", "text": f"[图片附件过大: {os.path.basename(path)}]"}
            with open(path, "rb") as f:
                b64 = base64.b64encode(f.read()).decode("ascii")
            ext = os.path.splitext(path)[1].lower().lstrip(".") or "png"
            if ext == "jpg":
                ext = "jpeg"
            return {"type": "image_url", "image_url": {"url": f"data:image/{ext};base64,{b64}"}}
        except Exception as e:
            return {"type": "text", "text": f"[图片编码失败: {e}]"}

    def toggle_voice(self):
        """语音输入：点击开始录音，再点击停止→识别→填入输入框"""
        if getattr(self, "voice_worker", None) and self.voice_worker.isRunning():
            # 正在录音 → 点击停止
            self.voice_worker.stop()
            self.set_status("识别中…")
            self.show_system_tip("录音已停止，正在识别…")
            return
        self.set_status("🎤 录音中…再次点击停止（最长12秒）")
        self.show_system_tip("🎤 录音中，点击「语音输入」按钮停止")
        self.voice_worker = VoiceRecWorker(self)
        self.voice_worker.done.connect(self._on_voice_done)
        self.voice_worker.error.connect(self._on_voice_error)
        self.voice_worker.start()

    def _on_voice_done(self, text):
        self.set_status("语音识别完成")
        self.input_box.insertPlainText(text + " ")
        self.input_box.setFocus()
        self.show_system_tip(f"🎤 已识别：{text}")

    def _on_voice_error(self, msg):
        self.set_status("语音输入失败")
        self.show_system_tip(f"语音输入失败：{msg}")

    # ---------- 拖拽文件 ----------
    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()

    def dropEvent(self, event):
        paths = [u.toLocalFile() for u in event.mimeData().urls() if u.isLocalFile()]
        if paths:
            self.add_pending_files(paths)

    def run_chat(self, text):
        if not getattr(self, "session_id", None):
            self.new_session()
        self.busy = True
        self._media_pending = []  # 新一轮对话重置媒体索引
        self.set_status("思考中…")
        self.show_user(text if text else "（查看附件）")
        user_content = text
        vision_blocks = []
        # 附件处理：文本类注入内容，图片类气泡预览 + 多模态视觉块
        if getattr(self, "pending_files", None):
            parts = []
            for p in self.pending_files:
                content, kind = self._read_attach(p)
                if kind == "text":
                    parts.append(f"【附件：{os.path.basename(p)}】\n{content}")
                    b = self._add_bubble("user")
                    b.set_text(f"📎 {os.path.basename(p)}")
                elif kind == "image":
                    b = self._add_bubble("user")
                    b.set_text(f"📷 {os.path.basename(p)}")
                    b.add_local_image(p)
                    parts.append(f"【附件图片：{os.path.basename(p)}，路径 {p}】")
                    vision_blocks.append(self._image_block(p))
                elif kind == "error":
                    self.show_system_tip(f"附件 {os.path.basename(p)} 读取失败：{content}")
                else:
                    self.show_system_tip(f"附件 {os.path.basename(p)} 格式暂不支持自动读取，可让 AI 尝试按路径处理")
                    parts.append(f"【附件：{os.path.basename(p)}（路径 {p}）】")
            if parts:
                user_content = (user_content + "\n\n" if user_content else "") + "\n\n".join(parts)
            self.pending_files = []
            self._refresh_attach_chips()
        if vision_blocks:
            content_parts = [{"type": "text", "text": user_content or "请分析这些图片"}] + vision_blocks
            self.messages.append({"role": "user", "content": content_parts})
        else:
            self.messages.append({"role": "user", "content": user_content})
        self._start_worker()

    def _start_worker(self):
        """构造消息并启动对话工作线程（供发送/重新回答复用）"""
        history = self.messages[-self.config.get("max_history", DEFAULT_CONFIG.get("max_history", 20)):]
        msgs = [{"role": "system", "content": self.config.get("system_prompt", DEFAULT_CONFIG["system_prompt"])}] + history
        tools = get_enabled_tools(getattr(self, "dept_state", None))
        self.worker = ChatWorker(self.clients, msgs, tools=tools, parent=self)
        self.worker.token.connect(self.on_token)
        self.worker.tool_used.connect(self.on_tool_used)
        self.worker.tool_result.connect(self.on_tool_result)
        self.worker.finished.connect(self.on_chat_finished)
        self.worker.error.connect(self.on_chat_error)
        self.worker.start()

    def regen_last_answer(self):
        """重新回答：移除最后一条 AI 回复（及其气泡），复用同一用户问题重新生成"""
        if self.busy or not getattr(self, "last_ai_bubble", None):
            return
        bubble = self.last_ai_bubble
        self._media_pending = []  # 重新生成时丢弃旧媒体索引
        # 移除 messages 中该轮 AI 回复（含工具调用中间消息）
        while self.messages and self.messages[-1].get("role") == "assistant":
            self.messages.pop()
        # 移除对应的气泡行
        row = getattr(bubble, "_row", None)
        if row is not None:
            self.chat_layout.removeWidget(row)
            row.deleteLater()
        self.last_ai_bubble = None
        self.streaming_active = False
        self.set_status("重新回答中…")
        self._start_worker()

    def on_token(self, tok):
        if not self.streaming_active:
            self.streaming_active = True
            self.last_ai_bubble = self._add_bubble("ai")
        self.last_ai_bubble.append_text(tok)
        self.set_status("正在生成…")

    def on_tool_used(self, msg):
        if msg.startswith("[工具]"):
            self.show_system_tip(msg)
        else:
            self.show_system_tip(msg)
        self.set_status(msg)

    def on_tool_result(self, result):
        kind = result.get("kind")
        if kind == "search":
            self.show_system_tip("")
            b = self._add_bubble("ai")
            b.add_search_card(result.get("query", ""), result.get("results", []))
        elif kind == "images":
            items = result.get("results", [])
            if not items:
                self.show_system_tip("没有找到相关图片")
                return
            self.show_system_tip(f"正在加载 {len(items)} 张图片…")
            self.img_worker = ImageDownloadWorker(items, self)
            self.img_worker.done.connect(self.on_images_downloaded)
            self.img_worker.start()
        elif kind == "image":
            path = result.get("path", "")
            msg = result.get("msg", "")
            if path and os.path.exists(path):
                self.show_system_tip("")
                b = self._add_bubble("ai")
                b.set_text(msg)
                b.add_local_image(path)
                self.session_files.add(path)
                self._media_pending.append({"type": "image", "path": path})
                self.save_current_session()
            else:
                self.show_system_tip(msg or "图片生成失败")
        elif kind == "video":
            path = result.get("path", "")
            msg = result.get("msg", "")
            if path and os.path.exists(path):
                self.show_system_tip("")
                b = self._add_bubble("ai")
                b.set_text(msg)
                b.add_video_card(path)
                self.session_files.add(path)
                self._media_pending.append({"type": "video", "path": path})
                self.save_current_session()
            else:
                self.show_system_tip(msg or "视频生成失败")

    def on_images_downloaded(self, saved):
        if saved:
            b = self._add_bubble("ai")
            b.add_image_grid(saved)
            for local, _, _, _ in saved:
                if os.path.isfile(local):
                    self.session_files.add(local)
                    self._media_pending.append({"type": "image", "path": local})
            self.save_current_session()
        else:
            self.show_system_tip("图片加载失败（网络受限或图片源不可达）")

    def on_chat_finished(self, full):
        self.streaming_active = False
        self.set_status("就绪")
        if full:
            msg = {"role": "assistant", "content": full}
            # 把本轮生成的媒体（图片/视频/语音）按顺序挂到该条回复上
            if getattr(self, "_media_pending", None):
                msg["media"] = list(self._media_pending)
                self._media_pending = []
            self.messages.append(msg)
            # 给本轮 AI 气泡添加「重新回答」按钮（仅对最终正文回复）
            if self.last_ai_bubble is not None:
                self.last_ai_bubble.add_regen_btn(self.regen_last_answer)
        else:
            self.show_ai("（已完成工具调用）")
        self.save_current_session()
        self.load_history_list()
        # 语音条回复
        if self.config.get("tts_enabled", True) and full:
            bubble = self.last_ai_bubble
            self.tts = TTSWorker(full[:500], self)
            self.tts.done.connect(lambda p, b=bubble: self._on_tts_done(p, b))
            self.tts.start()
        self.busy = False

    def _on_tts_done(self, mp3, bubble):
        if mp3 and bubble is not None:
            bubble.add_voice_bar(mp3)
            self.session_files.add(mp3)
            self._media_pending.append({"type": "audio", "path": mp3})
            self.save_current_session()

    def on_chat_error(self, msg):
        self.streaming_active = False
        self.set_status("出错")
        self.show_ai(f"发生错误：{msg}")
        self.busy = False

    def open_settings(self):
        dlg = SettingsDialog(self.config, self)
        if dlg.exec():
            self.config = dlg.config
            save_config(self.config)
            self.mode = self.config.get("mode", "cloud")
            self.clients = build_clients(self.config, self.mode)
            self._update_mode_btn()
            self.set_status("设置已保存")
            QMessageBox.information(self, "设置", "设置已保存并生效。备用 Provider 会在主 API 不可用时自动接管。")


def html_escape(text):
    return (text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                .replace('"', "&quot;").replace("'", "&#39;"))


APP_STYLE = """
QWidget { font-family: "Microsoft YaHei UI", "Microsoft YaHei", sans-serif; font-size: 13px; color: #1f2937; }
QMainWindow, QWidget#chatArea { background: #f5f7fa; }
#topbar { background: #ffffff; border-bottom: 1px solid #e8ecf1; }
#appTitle { font-size: 16px; font-weight: 700; color: #111827; }
#onlineDot { color: #22c55e; font-size: 12px; }
#statusLabel { color: #6b7280; font-size: 12px; }
#sidebar { background: #ffffff; border-right: 1px solid #e8ecf1; }
#sidebar QLabel { color: #6b7280; font-size: 12px; }
#hint { color: #b3b8c0; font-size: 11px; }
#chatContainer { background: #f5f7fa; }
#inputFrame { background: #ffffff; border-top: 1px solid #e8ecf1; }
QTextEdit#InputEdit { border: 1px solid #e5e9f0; border-radius: 14px; padding: 9px 14px; background: #f7f9fc; font-size: 14px; }
QTextEdit#InputEdit:focus { border-color: #2b5bff; background: #ffffff; }
QPushButton#primaryBtn { background: #2b5bff; color: #ffffff; border: none; border-radius: 8px; padding: 9px 24px; font-weight: 600; }
QPushButton#primaryBtn:hover { background: #1e4fd8; }
QPushButton#primaryBtn:pressed { background: #1740b3; }
QPushButton#primaryBtn:disabled { background: #c3cbe0; color: #ffffff; }
QPushButton#ghostBtn { background: transparent; border: 1px solid #d8dee9; border-radius: 8px; padding: 6px 14px; color: #374151; }
QPushButton#ghostBtn:hover { background: #f3f5f9; }
QPushButton#modeBtn { background: transparent; border: 1px solid #d8dee9; border-radius: 8px; padding: 6px 14px; color: #374151; font-size: 12px; }
QPushButton#modeBtn:hover { background: #f3f5f9; }
QPushButton#modeBtn:checked { background: #2b5bff; border-color: #2b5bff; color: #ffffff; }
QPushButton#chip { background: #f3f5f9; border: 1px solid #e5e9f0; border-radius: 16px; padding: 5px 14px; color: #374151; font-size: 12px; }
QPushButton#chip:hover { background: #e8ecf4; }
#deptTitle { font-size: 12px; font-weight: 700; color: #9aa3b2; padding: 2px 4px; letter-spacing: 1px; }
#deptBlock { background: #ffffff; border: 1px solid #edf0f5; border-radius: 10px; }
#deptBlock:hover { border-color: #dbe3f5; }
#deptName { font-size: 13px; font-weight: 600; color: #1f2937; }
#deptAction { background: #f3f5f9; border: none; border-radius: 6px; padding: 4px 10px; color: #374151; font-size: 12px; }
#deptAction:hover { background: #e8ecf4; color: #2b5bff; }
#deptAction:disabled { color: #c3cbe0; background: #f8f9fb; }
QListWidget#sessionList { background: transparent; border: none; }
QListWidget#sessionList::item { border-radius: 10px; margin: 1px 0; }
QListWidget#sessionList::item:hover { background: #f3f5f9; }
QListWidget#sessionList::item:selected { background: #eaf0ff; }
QWidget#sessionItem { border-radius: 10px; background: transparent; }
QWidget#sessionItem[selected="true"] { background: #eaf0ff; }
#sessionTitle { font-size: 14px; font-weight: 600; color: #111827; }
#sessionPreview { font-size: 12px; color: #9ca3af; }
#msgTime { color: #b0b6c0; font-size: 11px; }
QDialog { background: #f7f9fc; }
QDialog QLineEdit, QDialog QComboBox, QDialog QTextEdit { border: 1px solid #d8dee9; border-radius: 6px; padding: 6px 10px; background: #ffffff; }
#bubbleUser { background: #95ec69; border-top-left-radius: 16px; border-top-right-radius: 4px; border-bottom-left-radius: 16px; border-bottom-right-radius: 16px; }
#bubbleUser QTextBrowser { color: #111111; }
#bubbleUser QLabel { color: #111111; }
#bubbleAi { background: #ffffff; border: 1px solid #eef1f5; border-top-left-radius: 4px; border-top-right-radius: 16px; border-bottom-left-radius: 16px; border-bottom-right-radius: 16px; }
#bubbleAi QTextBrowser { color: #1f2937; }
#bubbleAi QLabel { color: #1f2937; }
#systemTip { color: #9ca3af; font-size: 11px; }
#voiceBar { background: #ffffff; border: 1px solid #e5e9f0; border-radius: 16px; max-width: 260px; }
#voiceBar:hover { background: #f3f5f9; }
#searchCard { background: #ffffff; border: 1px solid #e5e9f0; border-radius: 10px; }
QTableWidget { background: #ffffff; border: 1px solid #e5e9f0; border-radius: 6px; }
QHeaderView::section { background: #f3f5f9; border: none; padding: 6px; font-weight: 600; }
QScrollBar:vertical { background: transparent; width: 10px; }
QScrollBar::handle:vertical { background: #d5dae3; border-radius: 5px; min-height: 30px; }
QScrollBar::handle:vertical:hover { background: #b8bfcc; }
QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical { height: 0; }
QScrollArea { background: transparent; border: none; }
QListWidget::scroll-bar:vertical { width: 8px; }
"""



def main():
    ensure_dirs()
    app = QApplication(sys.argv)
    app.setStyle("Fusion")
    if FLUENT_OK:
        try:
            setTheme(Theme.LIGHT)
        except Exception as _te:
            pass
        try:
            setThemeColor("#2b5bff")
        except Exception as _tce:
            pass
    window = MainWindow()
    window.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
